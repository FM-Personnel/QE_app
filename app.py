import streamlit as st
import hashlib
import requests
import os
import pytz
import time
import unicodedata
import importlib.metadata
import re
import tempfile
import uuid
import base64
import logging
from pydantic import BaseModel, validator
from dotenv import load_dotenv
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.http import models
from qdrant_client.http.models import Filter, FieldCondition, MatchValue, MatchText
from typing import List, Optional, Union, Dict, Any, Literal, Set, Tuple
from datetime import datetime, timedelta
from collections import defaultdict
from types import SimpleNamespace
from PyPDF2 import PdfReader
from docx import Document
from urllib.parse import urlparse, parse_qs

try:
    import pdfplumber
    print("✅ pdfplumber est installé et fonctionnel.")
except ImportError:
    print("❌ pdfplumber n'est pas installé. Exécutez : pip install pdfplumber")
    raise

#################################################################
# 1. CLASSES, FONCTIONS UTILITAIRES, INITIALISATION DES VARIABLES
#################################################################

# --- Persistance de l'historique via session_state ---
def save_historique():
    # Limite le nombre d'entrées pour éviter la surcharge (ex: 50 dernières)
    if len(st.session_state.full_historique) > 50:
        # Supprime les entrées les plus anciennes
        sorted_entries = sorted(
            st.session_state.full_historique.items(),
            key=lambda x: x[1]["metadata"].get("timestamp", "")
        )
        st.session_state.full_historique = dict(sorted_entries[-50:])
    st.session_state.historique_cache = st.session_state.full_historique.copy()

def load_historique():
    if "historique_cache" in st.session_state:
        st.session_state.full_historique = st.session_state.historique_cache.copy()

# Pour permettre de renommer une collection
if 'show_rename_modal' not in st.session_state:
    st.session_state.show_rename_modal = False
if 'current_doc_to_rename' not in st.session_state:
    st.session_state.current_doc_to_rename = None

# Initialisation des états de certaines variables si non existants
if 'use_priority_docs' not in st.session_state:
    st.session_state.use_priority_docs = False
if 'priority_docs' not in st.session_state:
    st.session_state.priority_docs = []

# Estimation du nombre de tokens d'un texte en français
def estimate_tokens(text):
    """Estime le nombre de tokens pour Mistral Large (1 token ≈ 4 caractères en français)."""
    return len(text) // 4

# Fonction utilitaire pour tronquer le texte
def truncate_text(text: str, max_tokens: int = 500) -> str:
    """Tronque un texte à un nombre maximal de tokens (1 token ≈ 4 caractères)."""
    max_chars = max_tokens * 4
    return (text[:max_chars] + "...") if len(text) > max_chars else text

# Fonction de conversion des dates
def safe_parse_date(date_str: Optional[str]) -> datetime:
    """Convertit une date hétérogène en datetime, ou datetime.min si invalide."""
    if not date_str:
        return datetime.min
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%Y/%m/%d"):
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    return datetime.min

# Masquer les messages de warning de pdfminer
logging.getLogger("pdfminer").setLevel(logging.WARNING)

# Charger l'historique au démarrage
load_historique()

# --- Configuration de la page (doit être unique et en premier) ---
st.set_page_config(
    page_title="Parlement RAG",
    page_icon="🗳️",
    layout="wide"   # ← élargit toute la page
)

# Moteur de recherche internet : détail d'implémentation, plus exposé dans l'UI.
MOTEUR_RECHERCHE_INTERNET = "Google"  # "Google" | "Tavily"

# Modèle Mistral : libellé affiché -> taille passée à l'API ("mistral-<taille>-latest").
# Le sélecteur vit dans « Options avancées » (chemin parlementaire) ; ailleurs on
# retombe sur MODELE_MISTRAL_DEFAUT.
_MODEL_LABELS = {
    "Large (recommandé)": "large",
    "Medium": "medium",
    "Small (rapide, moins fiable)": "small",
}
MODELE_MISTRAL_DEFAUT = "large"

# --- Classe  CSS pour tous les messages de statut
st.markdown("""
<style>
    /* ===== STYLES GLOBAUX ===== */
    /* Messages de statut (utilisé par status_placeholder et prep_placeholder) */
    .status-message, .prep-message {
        display: flex;
        justify-content: center;
        align-items: center;
        margin: 0.5rem auto !important;
        padding: 0.7rem 1rem;
        text-align: center;
        font-size: 14px;
        color: #555;
        background-color: #f0f2f6;
        border-radius: 6px;
        max-width: 600px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    }
    /* Style spécifique pour les messages dans les onglets */
    .prep-message {
        min-height: 40px;
        font-size: 15px;
    }
    /* Conteneur principal */
    .stApp {
        display: flex;
        flex-direction: column;
    }
    .block-container {
        padding-top: 1rem;
        padding-bottom: 1rem;
    }
    /* Expanders et contenu */
    div[data-testid="stExpander"] {
        margin: 0 auto 0.5rem auto !important;
        max-width: 1000px;
        width: 100%;
    }
    div.streamlit-expanderHeader {
        padding: 0.3rem 0.6rem !important;
        font-size: 13px !important;
        background-color: #f0f2f6 !important;
        border-radius: 4px !important;
    }
    div.streamlit-expanderContent {
        padding: 0.5rem 0.8rem !important;
        text-align: justify;
    }
    /* Titres et séparateurs */
    h3 {
        text-align: center;
        margin: 0.4rem auto !important;
        color: #0066cc;
    }
    hr {
        margin: 0.3rem auto;
        width: 80%;
        border: none;
        border-top: 1px solid #ddd;
    }
    /* Suppression des marges inutiles */
    .stMarkdown > div {
        margin: 0 !important;
    }
    p, ul, ol {
        margin: 0.2rem 0 !important;
        padding: 0 !important;
    }
    /* Style pour les expanders de l'historique */
    div[data-testid="stExpander"] > details > summary {
        padding: 0.5rem 1rem !important;
        background-color: #f0f2f6 !important;
        border-radius: 6px !important;
        margin-bottom: 0.5rem !important;
    }
    /* Espacement entre les entrées */
    .history-entry {
        margin-bottom: 1rem !important;
    }
</style>
""", unsafe_allow_html=True)

# --- Chargement des variables d'environnement ---
load_dotenv()
QDRANT_URL = os.getenv("QDRANT_URL", "").strip()
QDRANT_API_KEY = os.getenv("QDRANT_API_KEY", "").strip()
MISTRAL_API_KEY = os.getenv("MISTRAL_API_KEY", "").strip()
TAVILY_API_KEY = os.getenv("TAVILY_API_KEY", "").strip()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY", "").strip()
GOOGLE_CX = os.getenv("GOOGLE_CX", "").strip()

QDRANT_COLLECTION = "QuestionParlementaire"

if not QDRANT_URL.startswith("https://"):
    raise RuntimeError("❌ QDRANT_URL doit commencer par https://")
if not QDRANT_API_KEY:
    raise RuntimeError("❌ QDRANT_API_KEY manquant. Vérifiez votre fichier .env")
if not MISTRAL_API_KEY:
    raise RuntimeError("❌ MISTRAL_API_KEY manquant. Vérifiez votre fichier .env")

print("QDRANT_URL:", QDRANT_URL)
print("QDRANT_API_KEY (début):", QDRANT_API_KEY[:10], "...")

# --- Chargement du modèle ---
LOCAL_MODEL_PATH = "./models/camembert_finetuned_progressive"
HUB_MODEL_PATH = "Whisler/camembert_finetuned_progressive"

@st.cache_resource
def load_embedding_model():
    try:
        if os.path.exists(LOCAL_MODEL_PATH):
            model = SentenceTransformer(LOCAL_MODEL_PATH)
            print("✅ Modèle chargé en local.")
        else:
            model = SentenceTransformer(HUB_MODEL_PATH)
            print("✅ Modèle téléchargé depuis HuggingFace Hub.")
        test_embedding = model.encode("Test de chargement du modèle.")
        VECTOR_SIZE = len(test_embedding)
        print("Dimension des embeddings:", VECTOR_SIZE)
        return model
    except Exception as e:
        st.error(f"❌ Erreur de chargement du modèle: {str(e)}")
        raise

embedding_model = load_embedding_model()

# --- Modèle BM25 sparse (recherche hybride sur la collection unifiée CodesJuridiques) ---
# fastembed "Qdrant/bm25" : tokenisation + IDF, léger (pas de réseau de neurones).
# Doit être appelé EXACTEMENT comme build_unified_codes.py (mêmes options par défaut)
# pour que le vecteur sparse de la requête soit comparable à ceux indexés.
@st.cache_resource
def load_bm25_model():
    try:
        from fastembed import SparseTextEmbedding
        m = SparseTextEmbedding("Qdrant/bm25")
        list(m.query_embed("test"))  # force le téléchargement / la mise en cache
        print("✅ Modèle BM25 (fastembed) chargé.")
        return m
    except Exception as e:  # noqa: BLE001
        print(f"⚠️ BM25 indisponible ({e}) — recherche juridique en mode vecteur seul.")
        return None

bm25_model = load_bm25_model()

# --- Connexion à Qdrant ---
try:
    qdrant_client = QdrantClient(
        url=QDRANT_URL,
        api_key=QDRANT_API_KEY,
        timeout=10.0,
        check_compatibility=False
    )
    collections = qdrant_client.get_collections()
    print(f"✅ Connexion réussie. Collections disponibles: {[c.name for c in collections.collections]}")
except Exception as e:
    print(f"❌ Erreur de connexion à Qdrant: {e}")
    raise


# --- Modèles Pydantic ---

# Modele Pydantic pour les articles juridiques
class BaseLegislativeRef(BaseModel):
    uid: str
    collection: str

class RetrievedLegalDocument(BaseModel):
    # Identifiants
    chunk_id: str
    num: str
    titre: str
    
    # Contenu
    contenu: str
    article_complet: str
    
    # Contexte
    contexte_hierarchique: str
    collection: str
    
    # Hiérarchie optionnelle
    partie: Optional[str] = None
    livre: Optional[str] = None
    titre_structure: Optional[str] = None
    chapitre: Optional[str] = None
    section: Optional[str] = None
    sous_section: Optional[str] = None
    paragraphe: Optional[str] = None
    sous_paragraphe: Optional[str] = None
    
    # Références législatives
    base_legislative: Optional[List[BaseLegislativeRef]] = None
    
    # Score du retrieval
    score: Optional[float] = None

# Modele Pydantic pour les documents generiques
class GenericDocument(BaseModel):
    # Identifiants
    uid: Optional[str] = None
    
    # Contenu
    text: str
    title: Optional[str] = None
    part: Optional[str] = None   # ex. "Annexe 7", "partie 1"
    
    # Métadonnées
    source: Optional[str] = None
    date_document: Optional[str] = None
    type_document: Optional[str] = None
    
    # Score du retrieval
    score: Optional[float] = None

# Modele Pydantic pour les reponses RAG
class ResponseDocument(BaseModel):
    # Identifiants
    uid: str
    # Contenu
    question: str
    reponse: str
    # Métadonnées
    legislature: Optional[str] = None
    chambre: Optional[str] = None   # Assemblée ou Sénat (à ajouter si tu l’as dans tes données)
    rubrique: Optional[str] = None
    analyse: Optional[str] = None
    ministeres_attribues: Optional[List[str]] = None
    # Dates
    date_question: Optional[str] = None
    date_reponse: Optional[str] = None
    # Références juridiques éventuelles
    textes_juridiques: Optional[List[str]] = None
    # Score du retrieval
    score: Optional[float] = None

#################################################################
# -------------- 1. PRINCIPALES FONCTIONS -----------------------
#################################################################

# --- 1a. Fonctions d'upload, d'indexation et d'embedding

# Fonction pour extraire le texte d'un document pdf
def extract_text(pdf_path: str, max_pages: Optional[int] = None) -> str:
    """
    Extrait le texte d'un PDF avec plusieurs stratégies :
    - pdfplumber pour le texte brut
    - fallback OCR si une page est vide
    - nettoyage des espaces et des sauts de ligne
    """
    text_chunks = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            pages = pdf.pages[:max_pages] if max_pages else pdf.pages
            for i, page in enumerate(pages, start=1):
                try:
                    # Extraction brute
                    page_text = page.extract_text() or ""
                    word_count = len(page_text.split()) if page_text else 0
                    
 #                    # Si vide, tenter une extraction par OCR (optionnel)
 #                    if not page_text.strip():
 #                        try:
 #                            from pdf2image import convert_from_path
 #                            import pytesseract
 #                            images = convert_from_path(pdf_path, first_page=i, last_page=i)
 #                            ocr_text = pytesseract.image_to_string(images[0], lang="fra")
 #                            page_text = ocr_text
 #                            st.write(f"Page {i} → OCR fallback → {len(page_text.split())} mots")
 #                        except Exception as e:
 #                            print(f"⚠️ OCR non disponible pour la page {i}: {e}")
 #                            pass
                    
                    # Nettoyage basique
                    page_text = re.sub(r"\s+", " ", page_text).strip()
                    
                    if page_text:
                        text_chunks.append(page_text)
                except Exception as e:
                    print(f"⚠️ Erreur page {i}: {e}")
                    continue
    except Exception as e:
        print(f"❌ Erreur lors de l'ouverture du PDF: {e}")
        return ""
    
    raw_text = "\n\n".join(text_chunks).strip()
    return raw_text

# Fonction pour extraire le texte d'un document Word
def extract_text_from_docx(file_path: str) -> str:
    """Extrait le texte d'un fichier Word."""
    try:
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return text.strip()
    except Exception as e:
        st.error(f"❌ Erreur lors de l'extraction du DOCX: {e}")
        return ""

# Fonction pour nettoyer les textes extraits
def clean_text(text: str) -> str:
    """Nettoie le texte extrait d'un PDF administratif en préservant la structure et les données utiles."""
    lines = text.split('\n')
    cleaned_lines = []

    for line in lines:
        line = line.strip()

        # Ignore les lignes vides ou presque vides
        if not line:
            continue

        # Ignore les numéros de page isolés (ex: "Page 57" ou "57")
        if re.match(r'^(?:Page\s*)?\d+\s*$', line, re.IGNORECASE):
            continue

        # Ignore les en-têtes/pieds de page répétitifs (ex: "PLFSS 2025 - Annexe 7")
        if re.match(r'^(?:PLFSS\s*\d{4}\s*-\s*Annexe\s*\d+|ANNEXE\s*DÉPENSES\s*DE\s*LA\s*BRANCHE\s*.*|securite-sociale\.fr|Source\s*:?\s*.*|Génération\s*X-Book)$', line, re.IGNORECASE):
            continue

        # Ignore les lignes avec seulement des caractères spéciaux ou des tirets
        if re.match(r'^[•○◘\-—~=]+$', line):
            continue

        # Conserve les lignes même courtes (titres, sous-titres, etc.)
        cleaned_lines.append(line)

    # Reconstitue le texte
    text = '\n'.join(cleaned_lines)

    # Nettoyage global
    text = re.sub(r'\s+', ' ', text)  # Espaces multiples
    text = re.sub(r'(\w)\s+-\s+(\w)', r'\1\2', text)  # Mots coupés par tiret
    text = re.sub(r'\bhttps?://\S+', '', text)  # URLs
    text = re.sub(r'[•○◘]+|[-=~]{5,}', '', text)  # Artefacts visuels

    # Nettoie les espaces résiduels
    text = text.strip()

    return text

# Fonction de suppression du sommaire
def remove_summary(text: str) -> str:
    """Supprime le sommaire et les tables des matières."""
    cleaned = re.sub(
        r"(?i)(SOMMAIRE|TABLE DES MATIÈRES).*?(?=PARTIE\s+\d+|ANNEXE\s+\d+|Article\s+\d+|$)",
        "",
        text,
        flags=re.DOTALL
    )
    return cleaned.strip() if cleaned.strip() else text

# Fonction de pré-traitement des titres
def preprocess_for_titles(text: str) -> str:
    """Insère des sauts de ligne avant chaque titre pour améliorer la segmentation."""
    sentences = re.split(r'(?<=[.!?])\s+', text)
    processed_text = []
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
        if detect_titles(sentence):
            processed_text.append(f"\n{sentence}\n")
        else:
            processed_text.append(sentence)
    text = " ".join(processed_text)

    # Ajout de sauts de ligne autour des titres
    title_patterns = [
        r"(Article\s*\d*\s*[-–]?)", r"(ANNEXE\s+\d+)", r"(PARTIE\s+\d+)",
        r"(Fiches?\s+d’?évaluation\s+préalable)", r"(\d+\.\s+)", r"([IVXLCDM]+\.\s+)"
    ]
    for pattern in title_patterns:
        text = re.sub(pattern, r"\n\1\n", text, flags=re.IGNORECASE)

    text = re.sub(r'\n\s*\n', '\n', text)
    return text.strip()

# Fonction de détection des titres
def detect_titles(line: str) -> bool:
    """Détecte les titres de manière robuste."""
    line = line.strip()
    if not line:
        return False
    regex_patterns = [
        r"^Article\s+\d+\s*[–—-]?", r"^ANNEXE\s+\d+", r"^PARTIE\s+\d+",
        r"^(TITRE|Chapitre|Section)\s+\d+", r"^[IVXLCDM]+\.\s+", r"^\d+(\.\d+)*\s+"
    ]
    if any(re.match(p, line, flags=re.IGNORECASE) for p in regex_patterns):
        return True
    title_keywords = [
        "Article ", "ANNEXE ", "PARTIE ", "Fiches d’évaluation préalable",
        "Synthèse", "Conclusion", "TITRE ", "Chapitre ", "Section ",
        "I. ", "II. ", "III. ", "1. ", "2. "
    ]
    return any(keyword.lower() in line.lower() for keyword in title_keywords)

# Fonction de segmentation du texte
def segment_text(text: str, max_words: int = 300, min_words: int = 50) -> List[Dict[str, str]]:
    """Découpe le texte en segments avec titres pour Qdrant."""
    blocks = re.split(r'\n\n|\.\s+', text)
    segments = []
    current_title = "AUTRE"
    current_content = []

    for block in blocks:
        block = block.strip()
        if not block:
            continue
        if detect_titles(block):
            if current_content:
                seg_text = ' '.join(current_content)
                if len(seg_text.split()) >= min_words:
                    segments.append({"title": current_title, "text": seg_text})
                current_content = []
            current_title = block
        else:
            current_content.append(block)

    if current_content:
        seg_text = ' '.join(current_content)
        if len(seg_text.split()) >= min_words:
            segments.append({"title": current_title, "text": seg_text})

    return segments

# Fonction pour préparer les chunks
def prepare_chunks_fixed(text: str, file_name: str, chunk_size=350, overlap=50) -> List[Dict]:
    """
    Découpe le texte en chunks fixes (~350 mots ≈ 512 tokens) avec overlap,
    en conservant le dernier titre détecté comme métadonnée.
    """
    words = text.split()
    chunks = []
    start = 0
    current_title = "AUTRE"

    while start < len(words):
        end = start + chunk_size
        chunk_words = words[start:end]

        # Met à jour le titre courant si un mot ressemble à un titre
        for w in chunk_words:
            if detect_titles(w):
                current_title = w

        if len(chunk_words) >= 50:  # filtre chunks trop courts
            chunk_text = " ".join(chunk_words)
            chunks.append({
                "id": str(uuid.uuid4()),
                "text": chunk_text,
                "metadata": {
                    "source": file_name,
                    "section": current_title,
                    "position": start,
                    "word_count": len(chunk_words),
                    "upload_date": datetime.now().isoformat()
                }
            })
        start += chunk_size - overlap

    return chunks

# Fonction pour l'upload et l'indexation avec logs détaillés
def process_and_index_document(file_path: str, file_type: str, collection_name: str,
                               qdrant_client=None, embedding_model=None,
                               progress_callback=None):
    """Traite et indexe un document dans SA PROPRE COLLECTION avec suivi de progression et logs détaillés."""
    try:
        # 1. Extraction du texte
        try:
            if file_type == "pdf":
                raw_text = extract_text(file_path)
            else:
                raw_text = extract_text_from_docx(file_path)

            if not raw_text:
                if progress_callback:
                    progress_callback(0, 100, "Échec : extraction du texte")
                return False

            if progress_callback:
                progress_callback(5, 100, "Extraction du texte terminée")

        except Exception as e:
            st.error(f"❌ Erreur extraction: {repr(e)}")
            return False

        # 2. Nettoyage et segmentation
        try:
            cleaned_text = clean_text(raw_text)

            segments = segment_text(cleaned_text) or [{"title": "Document", "text": cleaned_text}]

            if progress_callback:
                progress_callback(10, 100, "Nettoyage et segmentation terminés")
        except Exception as e:
            st.error(f"❌ Erreur nettoyage/segmentation: {repr(e)}")
            return False

        # 3. Pré-traitement titres
        try:
            preprocessed_text = preprocess_for_titles(cleaned_text)
        except Exception as e:
            st.error(f"❌ Erreur preprocess_for_titles: {repr(e)}")
            return False        

        # 4. Chunking
        try:
            chunks = prepare_chunks_fixed(
                preprocessed_text,
                file_name=collection_name,
                chunk_size=350,   # ≈ 512 tokens
                overlap=50
            )
        except Exception as e:
            st.error(f"❌ Erreur chunking: {repr(e)}")
            return False

        if progress_callback:
            progress_callback(30, 100, f"Préparation des chunks terminée ({len(chunks)} chunks)")

        # 5. Génération des embeddings (par petits lots)
        texts = [chunk["text"] for chunk in chunks]
        embeddings = []
        total_chunks = len(chunks)

        for i in range(0, total_chunks, 5):
            batch_texts = texts[i:i+5]
            try:
                batch_embeddings = embedding_model.encode(batch_texts).tolist()
                embeddings.extend(batch_embeddings)
            except Exception as e:
                st.error(f"❌ Erreur génération embeddings: {repr(e)}")
                if progress_callback:
                    progress_callback(0, 100, f"Erreur génération embeddings: {str(e)}")
                return False

            if progress_callback:
                current_chunk = min(i + 5, total_chunks)
                progress_callback(30 + int(30 * current_chunk / total_chunks),
                                  100,
                                  f"Génération des embeddings : {current_chunk}/{total_chunks}")

        # 6. Indexation dans Qdrant (par petits lots avec réessais)
        points = [
            models.PointStruct(
                id=chunk["id"],
                vector=embedding,
                payload={"text": chunk["text"], **chunk["metadata"]}
            )
            for chunk, embedding in zip(chunks, embeddings)
        ]

        batch_size = 10
        max_retries = 2
        retry_delay = 2

        for i in range(0, len(points), batch_size):
            batch = points[i:i + batch_size]
            success = False
            retry_count = 0

            while retry_count < max_retries and not success:
                try:
                    qdrant_client.upsert(
                        collection_name=collection_name,
                        points=batch,
                        wait=True,
                    )
                    success = True
                except Exception as e:
                    retry_count += 1
                    st.error(f"⚠️ Erreur upsert batch {i//batch_size+1}: {repr(e)}")
                    if progress_callback:
                        progress_callback(0, 100,
                                          f"Échec batch {i//batch_size + 1}, tentative {retry_count}/{max_retries}")
                    if retry_count < max_retries:
                        time.sleep(retry_delay)

            if not success:
                if progress_callback:
                    progress_callback(0, 100, f"Échec définitif batch {i//batch_size + 1}")
                return False

            if progress_callback:
                current_point = min(i + batch_size, len(points))
                progress_callback(60 + int(40 * current_point / len(points)),
                                  100,
                                  f"Indexation : {current_point}/{len(points)} chunks")

        if progress_callback:
            progress_callback(100, 100, "Indexation terminée avec succès")
        st.success("🎉 Document indexé avec succès")
        return True

    except Exception as e:
        if progress_callback:
            progress_callback(0, 100, f"Erreur : {str(e)}")
        st.error(f"Erreur dans process_and_index_document: {e}")
        return False

# --- 1b. Fonctions de recherche ---

# Fonction qui supprime les accents
def normalize_query(query: str) -> str:
    # Supprime les accents et normalise en ASCII
    return ''.join(
        c for c in unicodedata.normalize('NFD', query)
        if unicodedata.category(c) != 'Mn'
    )

# Fonction qui extrait le sujet principal de la question
def extract_subject(question: str) -> str:
    """
    Extrait le sujet principal d'une question parlementaire sous forme de 3 à 5 mots-clés.
    - Supprime toute mention de députés, du Gouvernement ou de formulations inutiles.
    - Nettoie la ponctuation et tronque à quelques mots.
    """
    url = "https://api.mistral.ai/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {MISTRAL_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "mistral-small-latest",  # possibilité de remplacer "small" par "medium"
        "messages": [
            {
                "role": "system",
                "content": (
                    "Tu es un assistant qui identifie uniquement le sujet principal "
                    "d'une question parlementaire. "
                    "Ne mentionne jamais le député ou le gouvernement. "
                    "Donne une réponse sous forme de 3 à 5 mots-clés concis, "
                    "centrés sur le thème (pas de phrase complète)."
                )
            },
            {
                "role": "user",
                "content": question
            }
        ],
        "temperature": 0.2,
        "max_tokens": 20
    }

    response = requests.post(url, headers=headers, json=payload)
    response.raise_for_status()
    data = response.json()

    try:
        subject = data["choices"][0]["message"]["content"].strip()
    except (KeyError, IndexError):
        subject = "sujet non identifié"

    # Nettoyage supplémentaire côté Python
    subject = re.sub(r"\b(députée?|député|gouvernement|ministre|assemblée nationale|sénat)\b", "", subject, flags=re.IGNORECASE)
    subject = subject.replace("**Sujet principal**", "").strip()
    subject = re.sub(r"[^\w\s]", " ", subject)  # supprime ponctuation
    subject = re.sub(r"\s+", " ", subject)

    # Tronquer à 5 mots max
    tokens = subject.split()
    subject = " ".join(tokens[:5])

    print("=== Sujet extrait (compact) ===", subject)

    return subject

# Fonction de recherche Tavily (et filtre pour les scores de pertinence <0.5)
def search_tavily_government(subject: str, min_score: float = 0.5):
    """
    Recherche les annonces gouvernementales récentes sur un sujet donné via Tavily.
    - Filtre par domaines autorisés
    - Filtre par date (moins d'un an si disponible)
    - Filtre par score de pertinence (>= min_score)
    """

    url = "https://api.tavily.com/search"
    headers = {"Authorization": f"Bearer {TAVILY_API_KEY}"}

    allowed_domains = [
        "gouvernement.fr", "education.gouv.fr", "vie-publique.fr", "elysee.fr",
        "solidarites.gouv.fr", "sante.gouv.fr", "travail-sante-solidarites.gouv.fr",
        "securite-sociale.fr", "ameli.fr", "lassuranceretraite.fr", "caf.fr",
        "msa.fr", "urssaf.fr", "legifrance.gouv.fr", "drees.solidarites-sante.gouv.fr",
        "ars.sante.fr", "cnsa.fr", "en3s.fr", "francetravail.fr"
    ]

    payload = {
        "query": f"dernières annonces gouvernement France {subject}",
        "max_results": 100,
        "include_answer": True,
        "include_domains": allowed_domains
    }

    response = requests.post(url, json=payload, headers=headers)
    response.raise_for_status()

    data = response.json()
    raw_results = data.get("results", [])

    # Filtre par domaine
    by_domain = [r for r in raw_results if any(d in r.get("url", "") for d in allowed_domains)]

    # Filtre par date (moins d'un an si dispo)
    cutoff = datetime.now() - timedelta(days=365)
    recent = []
    for r in by_domain:
        ds = r.get("published_date") or r.get("date")
        if ds:
            try:
                pub = datetime.fromisoformat(ds.replace("Z", ""))
                # ✅ Condition supplémentaire : année 2025
                if pub.year == 2025 or pub >= cutoff:
                    recent.append(r)
                    continue
                else:
                    continue
            except Exception:
                # date non exploitable → on garde
                recent.append(r)
        else:
            recent.append(r)

    # Filtre par score
    filtered = [r for r in recent if r.get("score", 0) >= min_score]
    filtered.sort(key=lambda x: x.get("score", 0), reverse=True)

    # ✅ Limiter à 10 après filtrage
    data["results"] = filtered[:10]

    return data

# Fonction de recherche Google
def search_google_government(subject: str,
                             min_score: float = 0.5,
                             max_results: int = 10):
    """
    Recherche des informations via Google Custom Search API sur un sujet donné.
    - Même structure et format que search_tavily_government
    - Entrée: subject (str), min_score, max_results
    - Sortie: dict {"results": [{"title","url","content","score","published_date"}]}
    """

    url = "https://www.googleapis.com/customsearch/v1"
    params = {
        "q": f"dernières annonces gouvernement France {subject}",
        "key": GOOGLE_API_KEY,
        "cx": GOOGLE_CX,
        "num": max_results
    }

    response = requests.get(url, params=params, timeout=20)
    response.raise_for_status()
    data = response.json()

    raw_results = data.get("items", [])
    allowed_domains = [
        "gouvernement.fr", "info.gouv.fr", "elysee.fr", "vie-publique.fr",
        "education.gouv.fr", "solidarites.gouv.fr", "sante.gouv.fr",
        "travail-sante-solidarites.gouv.fr", "securite-sociale.fr", "ameli.fr",
        "lassuranceretraite.fr", "caf.fr", "msa.fr", "urssaf.fr",
        "legifrance.gouv.fr", "drees.solidarites-sante.gouv.fr", "ars.sante.fr",
        "cnsa.fr", "en3s.fr", "francetravail.fr"
    ]

    # Filtre par domaine
    by_domain = [r for r in raw_results if any(d in r.get("link", "") for d in allowed_domains)]

    # Filtre par date (moins d'un an si dispo)
    cutoff = datetime.now() - timedelta(days=365)
    recent = []
    for r in by_domain:
        date_str = r.get("pagemap", {}).get("metatags", [{}])[0].get("article:published_time")
        if date_str:
            try:
                pub = datetime.fromisoformat(date_str.replace("Z", ""))
                if pub >= cutoff:
                    recent.append(r)
            except Exception:
                recent.append(r)
        else:
            recent.append(r)

    # Filtre par score (Google ne fournit pas de score → fallback = 1.0 si snippet présent)
    filtered = [r for r in recent if r.get("snippet")]
    filtered = filtered[:max_results]

    # Format homogène comme Tavily
    results = []
    for r in filtered:
        results.append({
            "title": r.get("title", ""),
            "url": r.get("link", ""),
            "content": r.get("snippet", ""),   # Tavily utilisait "content"
            "score": 1.0,                      # Valeur par défaut
            "published_date": date_str if r.get("pagemap", {}).get("metatags") else None
        })

    return {"results": results}

# --- Fonction de log pour le debug ---
def log_debug(title: str, data: Any, max_length: int = 500):
    """Affiche un log de debug dans un fichier."""
    import os
    from datetime import datetime

    # Chemin absolu vers le dossier de logs (dans le répertoire courant)
    log_dir = os.path.join(os.getcwd(), "logs")

    # Créer le dossier s'il n'existe pas
    if not os.path.exists(log_dir):
        os.makedirs(log_dir)
        print(f"📁 Dossier créé : {log_dir}")

    # Chemin absolu vers le fichier de log
    log_file = os.path.join(log_dir, "debug_logs.txt")

    try:
        with open(log_file, "a", encoding="utf-8") as f:
            f.write(f"\n--- {title} ---\n")
            f.write(f"Timestamp: {datetime.now().isoformat()}\n")
            if isinstance(data, dict):
                for k, v in data.items():
                    f.write(f"{k}: {str(v)[:max_length]}\n")
            elif isinstance(data, list):
                f.write(f"List of {len(data)} items:\n")
                for i, item in enumerate(data[:3]):
                    f.write(f"  Item {i}: {str(item)[:max_length]}\n")
            else:
                f.write(f"{str(data)[:max_length]}\n")
            f.write("---\n")

        # Confirmation dans le terminal
        print(f"✅ Log enregistré dans {log_file}: {title}")

    except Exception as e:
        print(f"❌ Erreur lors de l'écriture du log : {str(e)}")

# Fonctions utilitaires car le champ base_legislative peut contenir soit des objets BaseLegislativeRef, soit des dicts (selon la provenance des données)
def get_ref_uid(ref) -> str | None:
    return ref.uid if hasattr(ref, "uid") else ref.get("uid")

def get_ref_collection(ref, default_collection: str) -> str:
    return ref.collection if hasattr(ref, "collection") else ref.get("collection", default_collection)

# Préparer pour préparer le contenu en vue d'un export .txt
def build_export_content(response_data: dict, mode: str, include_legal_articles: bool = False) -> str:
    lines = []

    # Question
    question = response_data.get("question")
    if question:
        lines.append("❓ Question\n")
        lines.append(str(question))
        lines.append("\n\n")

    # En-tête (Réponse ou Analyse)
    if mode == "analyse":
        lines.append("🔎 Analyse juridique\n")
    else:
        lines.append("📜 Réponse\n")
    lines.append(str(response_data.get("response", "Pas de texte généré")))
    lines.append("\n\n")

    # Résumé (si présent)
    summary = response_data.get("summary")
    if summary:
        lines.append("📝 Résumé\n")
        lines.append(str(summary))
        lines.append("\n\n")

    # Résultats de recherche (si présents)
    search_results = response_data.get("search_results", [])
    search_summary = response_data.get("search_context") or response_data.get("answer")
    if search_summary:
        lines.append("🌐 Résumé de la recherche internet\n")
        lines.append(str(search_summary))
        lines.append("\n\n")
    if search_results:
        lines.append("🌐 Résultats de recherche\n")
        for idx, item in enumerate(search_results, start=1):
            titre = item.get("title", "Sans titre")
            url = item.get("url", "")
            extrait = item.get("content") or item.get("snippet") or ""
            score = item.get("score", "N/A")
            date = item.get("published_date", "N/A")
            lines.append(f"{idx}. {titre}\n")
            if url:
                lines.append(f"   Lien: {url}\n")
            if extrait:
                lines.append(f"   Extrait: {extrait}\n")
            lines.append(f"   Score: {score} | Date: {date}\n\n")

    # Anciennes QE (si présentes)
    lines.append("🏛️ Anciennes QE\n")
    for doc in response_data.get("similar_documents", []):
        score = f"{getattr(doc, 'score', 0):.2f}" if getattr(doc, "score", None) is not None else "N/A"
        chambre = getattr(doc, "chambre", "Inconnue")
        lines.append(f"- QE {doc.uid} ({chambre}) - Score : {score}\n")
        lines.append(f"  Question: {doc.question}\n")
        lines.append(f"  Réponse: {doc.reponse}\n\n")

    # Articles juridiques
    if mode == "analyse" or include_legal_articles:
        lines.append("⚖️ Articles juridiques\n")

        sources = response_data.get("sources", [])
        if not sources:
            lines.append("Aucun article juridique enregistré.\n")
        else:
            for art in sources:
                lines.append(f"--- Article {art.get('num','N/A')} ({art.get('collection','N/A')}) ---\n")
                lines.append(f"Titre: {art.get('titre','')}\n")
                lines.append(f"Texte: {art.get('contenu','')}\n")

                # Relations parents (articles cités)
                if art.get("parents"):
                    lines.append("Articles cités:\n")
                    for parent in art["parents"]:
                        lines.append(f"- {parent.get('num','N/A')}: {parent.get('titre','')}\n")

                # Relations enfants (référencé par)
                if art.get("enfants"):
                    lines.append("Référencé par:\n")
                    for enfant in art["enfants"]:
                        lines.append(f"- {enfant.get('num','N/A')}: {enfant.get('titre','')}\n")
                lines.append("\n")

    return "\n".join(str(x) for x in lines)

# =====================================================================
# Recherche juridique unifiée (CodesJuridiques + fusion native Qdrant)
# ---------------------------------------------------------------------
# `CodesJuridiques` = les 4 codes en UNE collection Qdrant, avec un vecteur
# dense (CamemBERT, réutilisé) ET un vecteur sparse BM25. Une seule requête
# `points/query` fusionne (DBSF) trois sous-requêtes :
#   1. dense (similarité cosinus) ;
#   2. sparse BM25 (recouvrement lexical — rattrape l'article précis noyé
#      parmi ses voisins de chapitre, mode de défaillance dominant du dense) ;
#   3. dense restreint aux articles dont la hiérarchie (section / chapitre)
#      matche le sujet de la question.
# Mesuré sur le jeu qe-eval (18 cas) : legal_hit@8 0.78 -> 0.89, article
# exact dans le top-8 0.35 -> 0.60. 1 requête au lieu de 4.
# Repli automatique sur la recherche 4-collections si BM25 ou la collection
# unifiée sont indisponibles.
# =====================================================================
_UNIFIED_CODE_COLLECTION = "CodesJuridiques"
_UNIFIED_HIER_FIELDS = ("section", "chapitre", "titre_structure", "sous_section")


def _is_legal_infra_collection(name: str) -> bool:
    """Collections d'infrastructure de la recherche juridique (vecteurs NOMMÉS
    `dense`/`bm25`, ou sparse seul) : `CodesJuridiques` et les compagnes
    `<Code>__bm25`. Elles ne sont PAS des « collections-documents » : un
    `qdrant_client.search()` à vecteur simple s'y solde par un 400 « Not existing
    vector name ». À exclure de `search_uploaded_documents` et des sélecteurs de
    documents de l'UI.
    """
    return name == _UNIFIED_CODE_COLLECTION or name.endswith("__bm25")

_QE_STOP = set("""
le la les un une des du de d au aux et ou a l en dans sur pour par que qui quoi dont ou
il elle ils elles on se sa son ses leur leurs ce cet cette ces est sont etre a ont
plus moins tres bien alors donc car ni or mais comme si aussi tout tous toute toutes
gouvernement gouvernementale ministre ministere delegue deleguee depute deputee
senateur senatrice parlementaire question ecrite orale attention appelle attire
interroge interpelle alerte souhaite savoir connaitre demande concernant sujet
monsieur madame mme mr etat france francais francaise afin egalement notamment ainsi
cadre situation situations mesure mesures dispositif dispositions difficulte
difficultes possibilite possibilites modalites modalite mise oeuvre place prise
charge relative relatif relatives suite face lors entre leurs quelles quels quelle
conditions condition reforme portant consequences consequence nombreux nombreuses
nouveau nouvelle nouvelles souvent notre nos votre vos majoritairement constitues
rencontrees rencontres actuellement recemment aujourd hui
""".split())

# Sigles du domaine -> forme longue (le code cite la forme longue, la QE le sigle)
_QE_ACRONYMS = {
    "rsa": "revenu de solidarite active",
    "aah": "allocation aux adultes handicapes",
    "apa": "allocation personnalisee d autonomie",
    "pch": "prestation de compensation du handicap",
    "ajpa": "allocation journaliere du proche aidant",
    "ajap": "allocation journaliere d accompagnement d une personne en fin de vie",
    "mdph": "maison departementale des personnes handicapees",
    "ars": "agence regionale de sante",
    "essms": "etablissements et services sociaux et medico sociaux",
    "esms": "etablissements et services medico sociaux",
    "ehpad": "hebergement pour personnes agees dependantes",
    "ase": "aide sociale a l enfance",
    "mna": "mineurs non accompagnes",
    "cna": "conference nationale de l autonomie",
    "cnsa": "caisse nationale de solidarite pour l autonomie",
    "crip": "recueil des informations preoccupantes",
    "aspa": "allocation de solidarite aux personnes agees",
    "cmi": "carte mobilite inclusion",
    "cti": "complement de traitement indiciaire remuneration medico social",
    "segur": "revalorisation salariale professionnels sanitaire social medico social",
}


def _qe_norm(s: str) -> str:
    s = unicodedata.normalize("NFD", s or "")
    s = "".join(c for c in s if unicodedata.category(c) != "Mn")
    return s.lower()


def _qe_subject(question: str) -> str:
    """Isole le sujet : une QE dit 'X interroge Y sur <SUJET>'. On coupe l'amorce."""
    q = _qe_norm(question)
    m = re.search(r"\b(?:sur|concernant|quant a|au sujet de|a propos de)\b", q[20:])
    return q[20 + m.start():] if m else q


def _qe_expand_acronyms(q: str) -> str:
    for sig, long in _QE_ACRONYMS.items():
        if re.search(rf"\b{sig}\b", q):
            q += " " + long
    return q


def _qe_hierarchy_phrases(question: str, limit: int = 8) -> list:
    """Bigrammes significatifs du sujet (pour le filtre plein texte sur la hiérarchie)."""
    q = _qe_expand_acronyms(_qe_subject(question))
    toks = [t for t in re.findall(r"[a-z0-9]+", q) if len(t) >= 4 and t not in _QE_STOP]
    bigrams, seen, out = [f"{a} {b}" for a, b in zip(toks, toks[1:])], set(), []
    for t in bigrams:
        if t not in seen:
            seen.add(t)
            out.append(t)
    return out[:limit]


def _qe_question_sparse(question: str) -> Optional[dict]:
    """Vecteur sparse BM25 du sujet de la question. None si BM25 indisponible."""
    if bm25_model is None:
        return None
    try:
        q = _qe_expand_acronyms(_qe_subject(question))
        e = list(bm25_model.query_embed(q))[0]
        return {"indices": [int(i) for i in e.indices], "values": [float(v) for v in e.values]}
    except Exception as e:  # noqa: BLE001
        print(f"⚠️ BM25 requête échouée ({e}) — repli vecteur seul.")
        return None


def unified_code_search(question: str, dense_vec: list, limit: int) -> Optional[list]:
    """Recherche juridique via `CodesJuridiques` + fusion native DBSF.

    Renvoie une liste d'objets normalisés (`.payload` dict + `.score`) prête pour
    `normalize_chunks`, ou `None` pour signaler qu'il faut retomber sur la
    recherche 4-collections (BM25 absent, collection absente, ou erreur réseau).
    """
    sparse = _qe_question_sparse(question)
    if sparse is None:
        return None  # sans BM25, l'unifié n'apporte rien (mesuré) -> repli

    prefetch = [
        {"query": dense_vec, "using": "dense", "limit": max(40, limit * 4)},
        {"query": sparse, "using": "bm25", "limit": max(40, limit * 4)},
    ]
    phrases = _qe_hierarchy_phrases(question)
    if phrases:
        should = [{"key": f, "match": {"text": p}}
                  for p in phrases for f in _UNIFIED_HIER_FIELDS]
        prefetch.append({"query": dense_vec, "using": "dense",
                         "limit": max(40, limit * 4), "filter": {"should": should}})

    body = {"prefetch": prefetch, "query": {"fusion": "dbsf"},
            "limit": max(30, limit * 6), "with_payload": True}
    try:
        r = requests.post(
            f"{QDRANT_URL.rstrip('/')}/collections/{_UNIFIED_CODE_COLLECTION}/points/query",
            json=body, headers={"api-key": QDRANT_API_KEY, "content-type": "application/json"},
            timeout=20)
        if r.status_code == 404:
            print(f"⚠️ Collection {_UNIFIED_CODE_COLLECTION} absente — repli 4-collections.")
            return None
        r.raise_for_status()
        points = r.json().get("result", {}).get("points", [])
    except Exception as e:  # noqa: BLE001
        print(f"⚠️ Recherche unifiée échouée ({e}) — repli 4-collections.")
        return None

    if not points:
        return None  # rien remonté : laisser la recherche 4-collections tenter sa chance

    return [SimpleNamespace(payload=p.get("payload", {}) or {}, score=p.get("score", 0.0))
            for p in points]


# --- F2 : quand la question cite explicitement un n° d'article, le forcer dans le
#     contexte (lookup direct) ; sinon, léger bonus lexical aux articles initiaux
#     dont le titre / la hiérarchie recoupe les termes saillants de la question.
_ARTICLE_CITE_RE = re.compile(
    r"\b([LRD])\.?\s?(\d{1,4}(?:\s?-\s?\d{1,4}){0,4})\b(?!\s?\d)"
)
_STOPWORDS_FR = {
    "les", "des", "une", "aux", "dans", "pour", "par", "sur", "avec", "sans", "sous",
    "que", "qui", "quoi", "dont", "cette", "ces", "leur", "leurs", "est", "sont",
    "être", "avoir", "plus", "moins", "entre", "vers", "chez", "afin", "lors",
    "cadre", "code", "article", "articles", "loi", "décret", "arrêté", "situation",
    "concernant", "relatif", "relative", "notamment", "ainsi", "cet",
}


def _norm_num(letter: str, digits: str) -> str:
    return letter.upper() + re.sub(r"\s", "", digits)


def extraire_articles_cites(text: str) -> list:
    """N° d'articles de code explicitement cités dans la question (['R4311-11-1', 'L4331-1'])."""
    out, seen = [], set()
    for m in _ARTICLE_CITE_RE.finditer(text or ""):
        num = _norm_num(m.group(1), m.group(2))
        if num not in seen and re.search(r"\d", num):
            seen.add(num)
            out.append(num)
    return out[:6]


def lookup_articles_par_num(nums: list, collections: list, debug: bool = False) -> list:
    """Récupère par filtre exact `num` chaque article cité (provenance='cited', score sentinelle haut)."""
    found = []
    for num in nums:
        hit = None
        for coll in collections:
            try:
                pts, _ = qdrant_client.scroll(
                    collection_name=coll,
                    scroll_filter=Filter(must=[FieldCondition(key="num", match=MatchValue(value=num))]),
                    limit=1, with_payload=True, with_vectors=False,
                )
                if pts:
                    hit = pts[0]
                    break
            except Exception:
                continue
        if hit is None:
            if debug:
                print(f"  article cité {num} : introuvable dans les codes")
            continue
        art = {**(hit.payload or {}), "provenance": "cited", "score": 3.0}
        art.setdefault("collection", coll)
        found.append(art)
    return found


def _tokens_saillants(text: str) -> set:
    toks = re.findall(r"[a-zàâäéèêëïîôöùûüç]{4,}", (text or "").lower())
    return {t for t in toks if t not in _STOPWORDS_FR}


def _bonus_lexical(article: dict, q_tokens: set) -> float:
    champ = f"{article.get('titre', '')} {article.get('contexte_hierarchique', '')}".lower()
    a_tokens = _tokens_saillants(champ)
    return min(0.4, 0.1 * len(q_tokens & a_tokens))


# Fonction qui recherche des articles juridiques dans les collections Qdrant en utilisant des embeddings
def search_articles(
    query: str,
    partie: Optional[str] = None,
    limit: int = 5,
    must_contain: Optional[str] = None,
    debug: bool = False,
    threshold: float = 0.0
) -> Dict[str, Any]:
    """Recherche optimisée d'articles juridiques dans Qdrant avec enrichissement (section + références)."""

    target_collections = ["CASF", "Code du travail", "Code de la santé publique", "Code de la sécurité sociale"]
    collections = qdrant_client.get_collections()
    valid_collections = [c.name for c in collections.collections if c.name in target_collections]

    if not valid_collections:
        return {"sources": [], "total": 0, "limit": limit, "offset": 0}

    try:
        query_filter = models.Filter(
            must=[models.FieldCondition(key="partie", match=models.MatchValue(value=partie))]
        ) if partie else None

        # --- 1. Extraction des sous-questions ---
        subqs = extract_subquestions(query)
        if debug:
            print("Sous-questions extraites :", subqs)

        fused_query = " ".join(subqs) if subqs else query
        embedding = embedding_model.encode(fused_query).tolist()

        # --- 1 bis. Recherche unifiée (CodesJuridiques + fusion DBSF dense/BM25/hiérarchie) ---
        # dense : vecteur des sous-questions fusionnées ; BM25 + filtre hiérarchie :
        # sujet ré-isolé depuis la question brute (l'heuristique attend le texte de la QE).
        unified_results = None if partie else unified_code_search(query, embedding, limit)

        if unified_results is not None:
            if debug:
                print(f"Recherche juridique : mode unifié ({len(unified_results)} candidats).")
            # les scores DBSF ne sont pas des cosinus -> le seuil cosinus ne s'applique pas ;
            # la fusion classe déjà par pertinence et seuls `limit` initiaux sont retenus (étape 6).
            results = unified_results
        else:
            all_results = []
            for collection in valid_collections:
                try:
                    hits = qdrant_client.query_points(
                        collection_name=collection,
                        query=embedding,
                        query_filter=query_filter,
                        limit=limit * 2,  # on prend plus large pour filtrer ensuite
                        with_payload=True,
                        with_vectors=False
                    ).points
                    all_results.extend(hits)
                except Exception:
                    continue

            # --- 2. Filtrage par score (cosinus) ---
            results = [r for r in all_results if r.score >= threshold]

        # --- 3. Filtrage par mot-clé ---
        if must_contain:
            results = [
                r for r in results
                if must_contain.lower() in r.payload.get("contenu", "").lower()
            ]

        if not results:
            return {"sources": [], "total": 0, "limit": limit, "offset": 0}

        # --- 4. Normalisation en articles initiaux ---
        initial_articles = normalize_chunks(results, provenance="initial")

        # --- 5. Déduplication + tri des initiaux ---
        seen = set()
        unique_initials = []
        for art in initial_articles:
            if art["num"] not in seen:
                seen.add(art["num"])
                unique_initials.append(art)

        # --- 5 bis. F2 : articles cités dans la question + bonus lexical ---
        articles_cites = extraire_articles_cites(query)
        if not articles_cites:
            # aucun n° cité : on reclasse les initiaux par recouvrement lexical
            # (titre + hiérarchie) avec les termes saillants de la question.
            q_tokens = _tokens_saillants(query)
            for art in unique_initials:
                base = art.get("score", 0) or 0
                art["score"] = base + _bonus_lexical(art, q_tokens)

        # Tri par score ou numéro
        unique_initials.sort(key=lambda d: d.get("score", 0), reverse=True)

        # --- 6. Limitation des initiaux ---
        limited_initials = unique_initials[:limit]

        # --- 7. Enrichissements ---
        context_articles = enrich_same_context(limited_initials, valid_collections)
        referenced_articles = enrich_references(limited_initials, valid_collections)

        # --- 8. Fusion + déduplication finale ---
        all_articles = limited_initials + context_articles + referenced_articles
        seen = set()
        final_articles = []
        for art in all_articles:
            if art["num"] not in seen:
                seen.add(art["num"])
                final_articles.append(art)

        # --- 8 bis. F2 : forcer la présence des articles explicitement cités ---
        if articles_cites:
            for art in final_articles:
                if art.get("num") in articles_cites:
                    art["provenance"] = "cited"
            manquants = [n for n in articles_cites if n not in seen]
            if manquants:
                forces = lookup_articles_par_num(manquants, valid_collections, debug=debug)
                if debug and forces:
                    print(f"Articles cités forcés dans le contexte : {[a.get('num') for a in forces]}")
                final_articles = forces + final_articles

        # --- 9. Tri final ---
        final_articles.sort(key=lambda d: (d.get("provenance") != "cited", str(d.get("num", ""))))

        return {
            "sources": final_articles,
            "total": len(final_articles),
            "limit": limit,
            "offset": 0
        }

    except Exception:
        import traceback
        traceback.print_exc()
        return {"sources": [], "total": 0, "limit": limit, "offset": 0}

# Fonction qui enrichit la liste des articles initiaux avec les articles de même niveau
def enrich_same_context(initial_articles: list, collections: list, debug=True) -> list:
    """
    Version finale corrigée utilisant uniquement scroll avec des filtres
    """
    enriched = []

    for art in initial_articles:
        if debug:
            print(f"\n🔹🔹🔹 Traitement de l'article {art.get('num', 'N/A')} 🔹🔹🔹")
            print(f"Collection: {art.get('collection', 'N/A')}")
            print(f"Contexte: {art.get('contexte_hierarchique', 'N/A')}")

        # Vérification des champs obligatoires
        if not all(k in art and art[k] for k in ['num', 'collection', 'contexte_hierarchique']):
            if debug:
                print("❌ Champs obligatoires manquants")
            continue

        try:
            # Utilisation de scroll avec filtre
            result = qdrant_client.scroll(
                collection_name=art['collection'],
                scroll_filter=Filter(
                    must=[
                        FieldCondition(
                            key="contexte_hierarchique",
                            match=MatchValue(value=art['contexte_hierarchique'])
                        )
                    ],
                    must_not=[
                        FieldCondition(
                            key="num",
                            match=MatchValue(value=art["num"])
                        )
                    ]
                ),
                limit=50,
                with_payload=True,
                with_vectors=False
            )

            points = result[0]  # La liste des points est le premier élément du tuple retourné

            if debug:
                print(f"Résultats: {len(points)} articles trouvés")
                for i, point in enumerate(points[:3]):  # Affiche les 3 premiers
                    print(f"  {i+1}. {point.payload.get('num', 'N/A')}")
                    print(f"     Titre: {point.payload.get('titre', 'N/A')}")
                    print(f"     Contexte: {point.payload.get('contexte_hierarchique', 'N/A')}")

            enriched.extend(normalize_chunks(points, provenance="context"))

        except Exception as e:
            if debug:
                print(f"❌ Erreur: {str(e)}")
                import traceback
                traceback.print_exc()
            continue

    if debug:
        print(f"\n📌 Résumé:")
        print(f"  Articles initiaux: {len(initial_articles)}")
        print(f"  Articles enrichis: {len(enriched)}")
        if enriched:
            print(f"  Numéros: {[a.get('num') for a in enriched[:5]]}")
        else:
            print("  Aucun article enrichi trouvé")

    return enriched

# Fonction pour enrichir la liste des articles avec les articles référencés
def enrich_references(initial_articles: list, collections: list, debug=True) -> list:
    """
    Version finale corrigée pour les références
    """
    enriched = []

    for art in initial_articles:
        if debug:
            print(f"\n🔗🔗🔗 Traitement des références pour {art.get('num', 'N/A')} 🔗🔗🔗")

        if 'base_legislative' not in art or not art.get('base_legislative'):
            if debug:
                print("⚠️ Pas de base_legislative")
            continue

        if debug:
            print(f"Base législative: {art['base_legislative']}")

        for ref in art['base_legislative']:
            ref_num = ref.get('uid')
            ref_collection = ref.get('collection', art['collection'])

            if not ref_num:
                if debug:
                    print(f"⚠️ UID manquant dans {ref}")
                continue

            if debug:
                print(f"Recherche de {ref_num} dans {ref_collection}")

            try:
                result = qdrant_client.scroll(
                    collection_name=ref_collection,
                    scroll_filter=Filter(
                        must=[
                            FieldCondition(
                                key="num",
                                match=MatchValue(value=ref_num)
                            )
                        ]
                    ),
                    limit=1,
                    with_payload=True,
                    with_vectors=False
                )

                points = result[0]

                if debug:
                    if points:
                        print(f"✅ Référence trouvée:")
                        for point in points:
                            print(f"    - {point.payload.get('num', 'N/A')}")
                    else:
                        print(f"❌ Référence {ref_num} non trouvée")

                enriched.extend(normalize_chunks(points, provenance="reference"))

            except Exception as e:
                if debug:
                    print(f"❌ Erreur: {str(e)}")
                    import traceback
                    traceback.print_exc()
                continue

    if debug:
        print(f"\n📌 Résumé des références:")
        print(f"  Références trouvées: {len(enriched)}")
        if enriched:
            print(f"  Numéros: {[a.get('num') for a in enriched[:5]]}")
        else:
            print("  Aucune référence trouvée")

    return enriched

# Fonction qui transforme les chunks en une liste d’articles normalisés
def normalize_chunks(hits: list, provenance: str) -> list:
    """Normalise les résultats Qdrant en conservant TOUS les champs."""
    return [{
        **hit.payload,  # Conserve tous les champs du payload
        "provenance": provenance,
        "score": hit.score if hasattr(hit, 'score') else None
    } for hit in hits]

# Fonction pour construire un arbre législatif
def build_legislative_tree(articles: list) -> dict:
    """
    Construit une arborescence législative à partir d'une liste d'articles normalisés (dicts).
    Niveaux utilisés : collection > partie > livre > titre_structure > chapitre > section > sous_section > paragraphe > sous_paragraphe > _items
    """
    def g(art: dict, key: str, default: str = "") -> str:
        # Accès sûr aux champs; tente aussi depuis 'payload' si présent
        if key in art and art[key] not in (None, ""):
            return art[key]
        payload = art.get("payload", {})
        val = payload.get(key)
        return val if val not in (None, "") else default

    def infer_partie_from_context(art: dict) -> str:
        ctx = g(art, "contexte_hierarchique")
        # Exemple de contexte: "Partie législative > Livre I : ... > Titre I : ..."
        if ctx:
            parts = [x.strip() for x in ctx.split(">")]
            for p in parts:
                if p.lower().startswith("partie"):
                    return p
        return g(art, "partie", "")

    def article_type_code(num: str) -> str:
        return num[0] if isinstance(num, str) and num else "?"

    tree: dict = {}

    for art in articles:
        collection = g(art, "collection")
        partie = infer_partie_from_context(art)
        livre = g(art, "livre")
        titre_structure = g(art, "titre_structure")
        chapitre = g(art, "chapitre")
        section = g(art, "section")
        sous_section = g(art, "sous_section")
        paragraphe = g(art, "paragraphe")
        sous_paragraphe = g(art, "sous_paragraphe")
        num = g(art, "num")
        titre = g(art, "titre")

        # Création des niveaux
        tree.setdefault(collection or "Collection inconnue", {})
        lvl1 = tree[collection or "Collection inconnue"]

        lvl1.setdefault(partie or "Partie inconnue", {})
        lvl2 = lvl1[partie or "Partie inconnue"]

        lvl2.setdefault(livre or "Livre inconnu", {})
        lvl3 = lvl2[livre or "Livre inconnu"]

        lvl3.setdefault(titre_structure or "Titre structure inconnu", {})
        lvl4 = lvl3[titre_structure or "Titre structure inconnu"]

        lvl4.setdefault(chapitre or "Chapitre inconnu", {})
        lvl5 = lvl4[chapitre or "Chapitre inconnu"]

        lvl5.setdefault(section or "Section inconnue", {})
        lvl6 = lvl5[section or "Section inconnue"]

        lvl6.setdefault(sous_section or "Sous-section inconnue", {})
        lvl7 = lvl6[sous_section or "Sous-section inconnue"]

        lvl7.setdefault(paragraphe or "Paragraphe inconnu", {})
        lvl8 = lvl7[paragraphe or "Paragraphe inconnu"]

        lvl8.setdefault(sous_paragraphe or "Sous-paragraphe inconnu", {})
        lvl9 = lvl8[sous_paragraphe or "Sous-paragraphe inconnu"]

        # Liste des items (articles)
        items = lvl9.setdefault("_items", [])

        # Métadonnées optionnelles (ex: type code L/R/D…)
        items.append({
            "num": num,
            "titre": titre,
            "type": article_type_code(num),
            "article": art  # garder l'article complet pour l'affichage détaillé
        })

    return tree

# Fonction de tri des numéros
def sort_key_article(article_data: Union[Dict[str, Any], str]) -> Tuple:
    """
    Clé de tri pour les numéros d'article.
    Gère les formats simples (L241-3) et complexes (D146-12-1).
    Retourne un tuple utilisable pour trier naturellement.
    """
    # Si on reçoit directement une chaîne
    if isinstance(article_data, str):
        num = article_data
    else:
        num = article_data.get("article", {}).get("num", "")

    if not isinstance(num, str) or not num:
        return ("", float("inf"))

    prefix = num[0]  # L, R, D...
    parts = re.findall(r"\d+", num)
    nums = [int(p) for p in parts] if parts else [float("inf")]

    return (prefix, *nums)

# Fonction qui ajoute un article dans l'arbre partagé selon les niveaux hiérarchiques
def add_to_tree(tree: dict, article: dict, item: dict):
    """
    Ajoute un article dans l'arbre hiérarchique avec ses relations.
    - article: dict normalisé (avec clés collection, partie, livre, etc.)
    - item: dict complet avec parents/enfants
    """
    # Construire la hiérarchie depuis les champs de l'article
    levels = [
        article.get("collection"),
        article.get("partie"),
        article.get("livre"),
        article.get("titre_structure"),
        article.get("chapitre"),
        article.get("section"),
        article.get("sous_section"),
        article.get("paragraphe"),
        article.get("sous_paragraphe"),
    ]
    # Filtrer les niveaux vides
    levels = [lvl for lvl in levels if lvl]

    # Naviguer dans l'arbre
    node = tree
    for label in levels:
        node = node.setdefault(label, {})

    # Ajouter l'article AVEC SES RELATIONS
    node.setdefault("_items", []).append(item)  # ✅ item contient parents/enfants

# Fonction qui affiche récursivement l'arbre sous forme d'expanders
def render_tree(container, node: dict, level: int = 0):
    """
    Affiche récursivement l'arbre sous forme d'expanders.
    Chaque article est un dict normalisé.
    """
    # Styles CSS
    container.markdown("""
    <style>
        .article-title { font-weight: bold; margin-bottom: 5px; }
        .article-body { font-size: 14px; line-height: 1.4; margin-bottom: 10px; }
        .relation-section { margin-top: 10px; margin-bottom: 10px; }
        .relation-title { font-weight: bold; color: #333; }
        .relation-item { margin-left: 15px; color: #555; }
    </style>
    """, unsafe_allow_html=True)

    # Affichage des articles avec tri basé sur sort_key_article
    for data in sorted(node.get("_items", []), key=sort_key_article):
        art = data["article"]

        with container.expander(f"📜 {art.get('num','?')} - {art.get('titre','')}"):
            # Contenu de l'article
            container.markdown(
                f'<div class="article-body">{art.get("contenu","")}</div>',
                unsafe_allow_html=True
            )

            # Section "Référencé par" (anciennement "Enfants")
            if data.get("enfants"):
                container.markdown(
                    '<div class="relation-section"><div class="relation-title">👶 Référencé par :</div>',
                    unsafe_allow_html=True
                )
                for enfant in data["enfants"]:
                    container.markdown(
                        f'<div class="relation-item">- {enfant.get("num","?")} : {enfant.get("titre","")}</div>',
                        unsafe_allow_html=True
                    )

            # Section "Articles cités" (anciennement "Parents")
            if data.get("parents"):
                container.markdown(
                    '<div class="relation-section"><div class="relation-title">📚 Articles cités :</div>',
                    unsafe_allow_html=True
                )
                for parent in data["parents"]:
                    container.markdown(
                        f'<div class="relation-item">- {parent.get("num","?")} : {parent.get("titre","")}</div>',
                        unsafe_allow_html=True
                    )

    # Navigation hiérarchique (uniquement pour l'organisation visuelle)
    for label, child in sorted(((k, v) for k, v in node.items() if k != "_items"), key=lambda x: x[0]):
        with container.expander(f"📁 {label}"):
            render_tree(container, child, level + 1)

# Fonction pour afficher les articles de façon plate (pour Réponse parlementaire)
def render_articles_flat(container, articles: list):
    """
    Affiche une liste plate d'articles normalisés (dicts) dans des expanders Streamlit.
    Chaque article est un dict avec les clés : num, titre, contenu, collection, contexte_hierarchique, score, provenance.
    """
    if not articles:
        container.info("Aucun texte juridique cité.")
        return

    for idx, art in enumerate(sorted(articles, key=lambda x: sort_key_article(x))):
        score = art.get("score", "N/A")
        label = f"{idx+1}. Article {art.get('num','?')} ({art.get('collection','N/A')}) - Score : {score}"

        with container.expander(label):
            container.markdown(f"**Titre :** {art.get('titre','')}")
            container.markdown(f"**Contexte hiérarchique :** {art.get('contexte_hierarchique','')}")
            container.markdown(f"**Texte complet :**\n\n{art.get('contenu','')}")

# Fonction de recherches de documents dans tout le RAG (hors codes et jeu de données QE) pour alimenter l'API
# --- F9 : requête distillée pour la recherche fiche ---
# Constat (05/09, cf. app_findings.md § F9) : le retrieval fiche est piloté par le
# recouvrement lexical dense ; une question réelle (1300-2700 caractères), chargée
# de références statutaires et de détails de l'espèce (dates, lieux, n° de décret),
# dilue le signal thématique dans le vecteur. Une sonde courte, topiquement dense,
# le concentre. `_qe_subject`/`_qe_expand_acronyms` (F2, déjà en prod) isolent déjà
# le sujet et développent les sigles pour rien ; on y ajoute les sous-questions
# reformulées (déjà calculées pour le bloc DEMANDES, aucun appel Mistral de plus).
# Débrayable à 4 valeurs. **MESURÉ le 06/09** (`feat/eval`, 4 runs
# `--fiche-vector` sur 6 cas, score brut de fiche) — les noms du harnais
# d'éval diffèrent des valeurs ci-dessous, correspondance vérifiée sur les
# 18 cas de `distilled-queries.json` (recalcul identique 18/18) :
#
#   valeur ici              harnais      clé d'embedding      score brut moyen
#   "off"                   raw          <id> (question)      0,708
#   "sujet"                 subject      <id>::fiche_sujet    0,749
#   "sujet_sousquestions"   distilled    <id>::fiche          0,759  ← DÉFAUT
#   "resume"                resume       <id>::fiche_resume   0,718
#
#   "off"                 : question brute (comportement d'avant F9).
#   "sujet"                : sujet isolé + sigles développés SEULS — court, dense,
#                            zéro appel Mistral de plus que la prod n'en fait déjà.
#   "sujet_sousquestions" : sujet + sous-questions reformulées (hypothèse F9
#                            initiale, 6712d41) — **DÉFAUT depuis le 06/09** :
#                            +0,052 sur la question brute, gagne sur les 5 cas
#                            où une fiche remonte, aucune régression. Aucun appel
#                            Mistral supplémentaire (les sous-questions sont déjà
#                            calculées pour le bloc DEMANDES).
#   "resume"               : résumé topique par UN appel Mistral dédié (voir
#                            resumer_requete_fiche) — mesuré 0,718, soit **à la
#                            fois le plus cher (+1 appel/génération) et moins bon
#                            que la distillation pure Python**. Conservé et testé,
#                            mais sans intérêt démontré ; la piste « v4 plus
#                            courte » est abandonnée.
#
# Réserves du verdict (détail dans app_findings.md § F9) : le gain est propre à
# l'écart texte-réel/fiche, quasi nul (+0,002) sur les sondes du reference-set
# déjà pré-alignées ; `qe-7021` reste non résolu et ne peut pas l'être par la
# distillation (écart de CONTENU de la fiche, pas de vocabulaire).
DISTILLER_REQUETE_FICHE = "sujet_sousquestions"  # "off" | "sujet" | "sujet_sousquestions" | "resume"


def distiller_requete_fiche(question: str, subquestions: Optional[List[str]] = None) -> str:
    """Requête dédiée à la recherche dans les collections `fiche_de_reference*`.

    Ne s'applique qu'à la recherche fiche (search_uploaded_documents) ; la
    recherche dans les autres collections-documents et tout le reste du
    retrieval (QE, articles) restent sur la question brute.
    """
    sujet = _qe_expand_acronyms(_qe_subject(question))
    if subquestions:
        return f"{sujet} {' '.join(subquestions)}".strip()
    return sujet


# Prompt exact (mandat Coordination, nuit du 05/09) — cité tel quel dans
# app_findings.md § F9 pour reproductibilité. Ne pas paraphraser sans mettre
# les deux à jour ensemble.
_RESUME_FICHE_SYSTEME = (
    "Vous résumez une question parlementaire en 2 à 4 phrases, dans le vocabulaire "
    "administratif standard d'une fiche de référence documentaire (pas le registre "
    "d'une question écrite). Consignes strictes : n'incluez aucune amorce "
    "parlementaire (ne mentionnez ni le nom du parlementaire, ni « appelle "
    "l'attention », ni le ministre) ; nommez le dispositif ou la mesure concernée ; "
    "nommez le ou les publics concernés ; citez les textes de loi, décrets ou "
    "arrêtés mentionnés avec leur numéro et leur date s'ils figurent dans la "
    "question ; employez les termes techniques du domaine (sigles développés une "
    "fois si utile). Rédigez en PROSE CONTINUE UNIQUEMENT : aucun markdown, aucun "
    "astérisque, aucun tiret, aucun titre, aucune liste à puces — des phrases "
    "complètes reliées entre elles, comme dans le corps d'une fiche de référence. "
    "Longueur impérative : 250 à 400 caractères MAXIMUM, pas plus. N'inventez aucun "
    "fait, aucune date, aucun numéro qui ne figure pas dans le texte fourni."
)


def resumer_requete_fiche(question: str, model: str = "mistral-small-latest") -> str:
    """F9 (mode "resume") — résumé topique de la question par un appel Mistral
    dédié, `temperature=0` (reproductibilité). Un appel supplémentaire PAR
    GÉNÉRATION quand ce mode est actif (latence/coût estimés dans
    PASSATION-app.md) ; aucun appel si le mode n'est pas "resume" (défaut
    inchangé). Les erreurs réseau/API remontent à l'appelant, qui retombe sur
    `fiche_query=None` (comportement identique aux autres modes, cf.
    generate_response).

    Même défaut que F10 possible ici : avec plusieurs références légales à
    citer, le modèle dépasse parfois le plafond de tokens et coupe en plein
    mot (« ...loi n° 86-33 du »). Filet de sécurité identique à F10 : si
    `finish_reason == "length"` et que le texte ne se termine pas proprement,
    on le ramène à sa dernière phrase complète plutôt que de renvoyer un
    fragment cassé (une référence tronquée serait pire qu'une référence en
    moins pour une requête d'embedding).
    """
    r = requests.post(
        "https://api.mistral.ai/v1/chat/completions",
        headers={"Authorization": f"Bearer {MISTRAL_API_KEY}", "Content-Type": "application/json"},
        json={
            "model": model,
            "messages": [
                {"role": "system", "content": _RESUME_FICHE_SYSTEME},
                {"role": "user", "content": question},
            ],
            "temperature": 0.0,
            "max_tokens": 160,  # cohérent avec les 400 c max demandés, marge incluse
        },
        timeout=30,
    )
    r.raise_for_status()
    data = r.json()
    resume = data["choices"][0]["message"]["content"].strip()
    if data["choices"][0].get("finish_reason") == "length" and not _se_termine_proprement(resume):
        dernier_point = resume.rfind(".")
        resume = resume[:dernier_point + 1].strip() if dernier_point > 0 else ""
    return resume


# --- F12 : règle de rétention des extraits documentaires, UNE SEULE définition ---
# Le mécanisme « fiche de référence » (contenu curé, chiffres datés) a trois
# paramètres, historiquement recopiés à la main à chaque point d'usage :
# un boost de score, un seuil standard, un seuil assoupli réservé aux fiches.
# Cette duplication a produit le bug F12 : l'onglet « Base documentaire »
# filtrait à 0,70 SANS la clause fiche, alors que l'injection dans le prompt
# l'appliquait — un extrait de fiche entre 0,45 et 0,70 partait au modèle
# sans jamais s'afficher, rendant l'onglet trompeur comme instrument
# d'observation. Tout passe désormais par ces constantes et `extrait_retenu`.
BOOST_FICHE = 0.12          # bonus de score appliqué aux extraits de fiches de référence
SEUIL_DOC_STANDARD = 0.70   # score minimal d'un extrait de document pour être retenu
SEUIL_FICHE_ASSOUPLI = 0.45 # seuil abaissé, réservé aux fiches de référence


def extrait_retenu(doc: Dict, min_score: float = SEUIL_DOC_STANDARD) -> bool:
    """Un extrait documentaire est-il retenu ? (F12 — règle unique)

    Retenu s'il atteint le seuil standard, OU s'il s'agit d'une fiche de
    référence atteignant le seuil assoupli. Utilisée par l'injection dans le
    prompt (`format_uploaded_docs_by_relevance`) ET par l'affichage de
    l'onglet « Base documentaire » : les deux doivent montrer le même
    ensemble, sinon l'onglet ment sur ce que le modèle a réellement reçu.
    """
    score = doc.get("score") or 0
    return score >= min_score or (bool(doc.get("is_fiche")) and score >= SEUIL_FICHE_ASSOUPLI)


def extrait_marginal(doc: Dict, min_score: float = SEUIL_DOC_STANDARD) -> bool:
    """Vrai si l'extrait n'est retenu QUE par la clause fiche assouplie —
    c'est-à-dire sous le seuil standard. Sert à le signaler à l'affichage,
    pour qu'un extrait faible ne se lise pas comme un extrait fort."""
    score = doc.get("score") or 0
    return bool(doc.get("is_fiche")) and SEUIL_FICHE_ASSOUPLI <= score < min_score


def search_uploaded_documents(
    query: str,
    qdrant_client: Any,
    embedding_model: Any,
    selected_collections: List[str] = None,
    top_k: int = 5,
    top_k_selected = 10,
    fiche_query: Optional[str] = None,
) -> List[Dict]:
    """
    Recherche dans les collections-documents (1 collection = 1 document).
    Args:
        query: Requête de recherche.
        qdrant_client: Client Qdrant.
        embedding_model: Modèle d'embedding.
        selected_collections: Liste des collections à rechercher (optionnel).
        top_k: Nombre de résultats max si aucune collection sélectionnée.
        top_k_selected: Nombre de résultats max si collections sélectionnées.
        fiche_query: si fourni et différent de `query`, requête utilisée pour
            EMBEDDER la recherche dans les collections fiche uniquement (F9).
    Returns:
        Liste de dicts avec les champs essentiels.
    """
    query_embedding = embedding_model.encode(query).tolist()
    # F9 : vecteur dédié aux fiches, un seul appel supplémentaire au modèle déjà
    # chargé (aucun nouveau modèle, aucun appel réseau) — seulement si la requête
    # distillée diffère de la question brute.
    fiche_embedding = (embedding_model.encode(fiche_query).tolist()
                       if fiche_query and fiche_query != query else query_embedding)
    all_results = []
    try:
        # 1. Récupère les collections à rechercher
        protected = {
            "QuestionParlementaire",
            "Code de la sécurité sociale",
            "Code du travail",
            "CASF",
            "Code de la santé publique",
        }
        collections = qdrant_client.get_collections()
        doc_collections = [col.name for col in collections.collections
                           if col.name not in protected
                           and not _is_legal_infra_collection(col.name)]

        # 2. Limite aux collections sélectionnées si spécifiées
        if selected_collections:
            doc_collections = [col for col in doc_collections if col in selected_collections]
            # le nombre de chunks retournés passe à 10 si la recherche est limité sur une ou plusieurs collections
            top_k = top_k_selected
            if not doc_collections:
                doc_collections = [col for col in doc_collections if col not in protected]
            
        # 3. Recherche dans chaque collection-document
        for collection in doc_collections:
            try:
                _is_fiche_coll = collection.lower().startswith("fiche_de_reference")
                results = qdrant_client.query_points(
                    collection_name=collection,
                    query=fiche_embedding if _is_fiche_coll else query_embedding,  # F9
                    limit=30 if _is_fiche_coll else top_k,  # fiche : on veut voir tous ses chunks
                    with_payload=True,
                    with_vectors=False,
                ).points
                # Bonus pour les "fiches de référence" curées (chiffres-clés datés,
                # textes récents) : elles doivent primer sur un chunk de PAP ou de
                # rapport ancien pour un chiffre, une date ou un sigle. Le modèle
                # d'embedding est entraîné sur des QE, pas sur des fiches -> leurs
                # scores bruts sont bas, d'où un bonus franc + une garantie de
                # présence (cf. point 4).
                is_fiche = _is_fiche_coll
                for result in results:
                    payload = result.payload or {}
                    score = float(result.score) if hasattr(result, "score") else None
                    if score is not None and is_fiche:
                        score = min(1.0, score + BOOST_FICHE)
                    all_results.append({
                        "collection": collection,
                        "text": payload.get("text", ""),
                        "score": score,
                        "title": payload.get("section"),
                        "is_fiche": is_fiche,
                    })
            except Exception as e:
                st.warning(f"Erreur sur {collection}: {e}")

        # 4. Trie par score et limite les résultats ; on garde EN PLUS tous les
        #    extraits de fiches de référence qui dépassent le seuil assoupli
        #    (contenu curé, court et peu nombreux : on préfère tout injecter que
        #    laisser un chiffre à jour hors du prompt). Plafond de sécurité : 8.
        all_results.sort(key=lambda x: (x["score"] is not None, x["score"]), reverse=True)
        top = all_results[:top_k]
        fiches = [r for r in all_results
                  if r.get("is_fiche") and (r.get("score") or 0) >= SEUIL_FICHE_ASSOUPLI][:8]
        for r in fiches:
            if r not in top:
                top.append(r)
        return top

    except Exception as e:
        st.error(f"Erreur de recherche dans les documents uploadés: {e}")
        return []

# Fonction qui ajuste la taille du contexte issu des "autres collections" à la pertinence (score) du retrieval
def format_uploaded_docs_by_relevance(
    uploaded_results: List[Dict],
    min_score: float = SEUIL_DOC_STANDARD,
    max_docs: int = 5,
    max_length: int = 400
) -> str:
    """
    Trie les résultats par score et formate les extraits textuels pour enrichir le prompt.
    - uploaded_results : liste de dictionnaires renvoyés par la recherche vectorielle
    - min_score : seuil de pertinence minimum
    - max_docs : nombre maximum de documents à inclure
    - max_length : longueur maximale de l'extrait
    """
    if not uploaded_results:
        return ""

    # Filtrer par score — règle unique `extrait_retenu` (F12), partagée avec
    # l'affichage de l'onglet « Base documentaire » pour que les deux montrent
    # le même ensemble. Seuil assoupli pour les fiches de référence (contenu
    # curé et daté, à faire remonter même quand le score brut reste modeste).
    filtered = [doc for doc in uploaded_results if extrait_retenu(doc, min_score)]

    # Trier par score décroissant
    filtered.sort(key=lambda d: d.get("score", 0), reverse=True)

    # Limiter le nombre de docs
    filtered = filtered[:max_docs]

    formatted = []
    for doc in filtered:
        passage = doc.get("text", "")
        source = doc.get("collection", "inconnu").replace("_", " ")
        title = doc.get("title", "") or ""
        score = doc.get("score", 0)
        # Les fiches de référence sont denses en chiffres datés : on leur laisse un
        # extrait plus long pour que la donnée utile ne soit pas coupée.
        limit = 1400 if doc.get("is_fiche") else max_length

        formatted.append(
            f"Source: {source} (Section: {title})\n"
            f"Passage: {passage[:limit]}...\n"
            f"(Score: {score:.2f})\n"
        )

    return "\n---\n".join(formatted)

# Fonction de recherches d'anciennes questions / réponses dans le RAG Qdrant
def search_question_parlementaire(query: str, top_k: int = 5) -> List[ResponseDocument]:
    embedding = embedding_model.encode(query).tolist()
    hits = qdrant_client.query_points(
        collection_name="QuestionParlementaire",
        query=embedding,
        limit=top_k,
        with_payload=True,
        with_vectors=False
    ).points
    results: List[ResponseDocument] = []

    for i, r in enumerate(hits):
        p = r.payload or {}
        try:
            # Normalisation sécurisée des champs
            uid = str(p.get("uid", "")) if p.get("uid") is not None else ""
            legislature = str(p.get("legislature")) if p.get("legislature") is not None else None
            ministeres = p.get("ministeres_attribues")
            if isinstance(ministeres, str):
                ministeres = [ministeres]
            elif not isinstance(ministeres, list):
                ministeres = []
            textes_juridiques = p.get("textes_juridiques")
            if isinstance(textes_juridiques, str):
                textes_juridiques = [textes_juridiques]
            elif not isinstance(textes_juridiques, list):
                textes_juridiques = []

            # Création du document
            doc = ResponseDocument(
                uid=uid,
                question=p.get("question", ""),
                reponse=p.get("reponse", ""),
                legislature=legislature,
                chambre=p.get("chambre"),
                rubrique=p.get("rubrique"),
                analyse=p.get("analyse"),
                ministeres_attribues=ministeres,
                date_question=p.get("date_question"),
                date_reponse=p.get("date_reponse"),
                textes_juridiques=textes_juridiques,
                score=r.score
            )
            results.append(doc)

        except Exception as e:
             continue  # Passe au résultat suivant

    return results

# Repondération du contexte parlementaire par la récence.
# Une réponse ministérielle de plusieurs années décrit un état du droit dépassé
# (plans périmés, chiffres anciens). Le tri par similarité pure fait parfois
# remonter une QE très proche mais ancienne, que le modèle recopie telle quelle.
# On garde le tri par similarité pour l'affichage ("Anciennes QE") ; pour le
# contexte injecté dans le prompt, on repondère : poids = similarité * facteur
# de récence à décroissance exponentielle (demi-vie ci-dessous).
RECENCY_HALF_LIFE_YEARS = 4.0

# Législature en cours : sert à distinguer une QE antérieure « du même dossier »
# (trame factuelle recoupable) d'une QE d'une mandature révolue (registre seul).
LEGISLATURE_COURANTE = "17"
_PROX_TRAME_MIN = 0.85          # proximité cosinus au-dessus de laquelle la trame est admissible
_TRAME_AGE_MAX_ANNEES = 2.0     # et réponse de moins de 2 ans


def _legislature_de(doc) -> Optional[str]:
    leg = getattr(doc, "legislature", None)
    if leg:
        return str(leg).strip()
    m = re.search(r"L(\d{1,2})", getattr(doc, "uid", "") or "")
    return m.group(1) if m else None


def annoter_qe_contexte(doc, now: Optional[datetime] = None) -> str:
    """En-tête d'un bloc de contexte parlementaire : proximité + statut trame/registre (F1)."""
    now = now or datetime.now()
    score = getattr(doc, "score", None)
    d = safe_parse_date(getattr(doc, "date_reponse", None))
    age = None if d == datetime.min else max(0.0, (now - d).days / 365.25)
    meme_leg = _legislature_de(doc) == LEGISLATURE_COURANTE
    recente = age is not None and age <= _TRAME_AGE_MAX_ANNEES
    admissible = (score or 0) >= _PROX_TRAME_MIN and meme_leg and recente
    statut = "TRAME FACTUELLE ADMISSIBLE" if admissible else "registre seulement"
    prox = f"proximité {score:.2f}" if isinstance(score, (int, float)) else "proximité n. d."
    return (f"[{prox} · {statut}] (source: {getattr(doc, 'uid', '?')}, "
            f"réponse du {getattr(doc, 'date_reponse', None) or 'date inconnue'})")


# --- Recherche internet : bloc daté / sourcé + repérage des textes publiés (F1) ---
# Avant : seuls les extraits bruts étaient injectés (ni titre, ni date, ni site),
# le modèle ne pouvait pas voir qu'un résultat « JORF … Décret n° … » attestait
# d'une parution — et écrivait « le Gouvernement travaille à l'élaboration d'un
# décret » (cas #6491). Chaque résultat porte maintenant [date · site] + titre,
# et un marqueur « TEXTE PUBLIÉ » quand le résultat donne un texte comme paru.
_REF_TEXTE_RE = re.compile(
    r"(?:d[ée]cret|arr[êe]t[ée]|loi|ordonnance)\s+n[°º]\s*\d{4}-\d+"
    r"(?:\s+du\s+\d{1,2}(?:er)?\s+[a-zéû]+\s+\d{4})?",
    re.IGNORECASE,
)
_PARUTION_RE = re.compile(r"JORF|Journal officiel|\bJO\s+(?:du|n[°º])", re.IGNORECASE)
_A_VENIR_RE = re.compile(
    r"projet de (?:d[ée]cret|loi|arr[êe]t[ée])|en pr[ée]paration|en cours d'[ée]laboration"
    r"|consultation publique|sera publi[ée]",
    re.IGNORECASE,
)


def detecter_texte_publie(titre: str, extrait: str) -> str:
    """Référence du texte si le résultat atteste d'une parution, sinon ''.

    Parution = mention JORF / Journal officiel, ou numéro attribué (« décret
    n° 2025-897 du 4 septembre 2025 ») — un projet ou une consultation sans
    numéro ne compte pas.
    """
    texte = f"{titre or ''} {extrait or ''}"
    ref = _REF_TEXTE_RE.search(texte)
    parution = _PARUTION_RE.search(texte)
    if not ref and not parution:
        return ""
    if not ref and _A_VENIR_RE.search(texte):
        return ""
    return " ".join(ref.group(0).split()) if ref else parution.group(0)


def _domaine_de(url: str) -> str:
    m = re.match(r"https?://(?:www\.)?([^/]+)", url or "")
    return m.group(1) if m else ""


def _date_courte(date_str) -> str:
    d = safe_parse_date(str(date_str)[:10]) if date_str else datetime.min
    return d.strftime("%d/%m/%Y") if d != datetime.min else ""


def annoter_resultat_internet(item: dict) -> str:
    """En-tête d'un résultat : [date · site] [TEXTE PUBLIÉ · réf] titre."""
    titre = (item.get("title") or "").strip()
    extrait = item.get("content") or item.get("snippet") or ""
    entete = " · ".join(b for b in (_date_courte(item.get("published_date") or item.get("date")),
                                     _domaine_de(item.get("url") or item.get("link"))) if b)
    ref = detecter_texte_publie(titre, extrait)
    parts = [f"[{entete}]" if entete else "", f"[TEXTE PUBLIÉ · {ref}]" if ref else "", titre]
    return " ".join(p for p in parts if p)


def _cle_extrait(item: dict) -> str:
    """Clé de dédoublonnage d'un résultat : extrait normalisé (160 premiers caractères)."""
    extrait = (item.get("content") or item.get("snippet") or "").lower()
    return re.sub(r"[^a-z0-9àâäéèêëïîôöùûüç]+", " ", extrait).strip()[:160]


# F3 — plafond des résultats internet, RÉFUTÉ par la mesure (06/09, voir app_findings.md).
# Contexte (nuit du 05/09) : survivants observés #6491→9, #8922→10, #3171→7,
# #5621→8 ; sur #8922, ~2-3 des 10 retenus étaient pertinents, le reste du bruit
# historique. Piste 1 (plafond absolu) appliquée avec un tri par recouvrement
# lexical pour ne pas découper arbitrairement (Google ne renvoie pas de score
# réel). MESURÉE le 06/09 sur #5621 dans l'app en ligne : le seul résultat
# pertinent-et-exploité des rangs 6-10 est au RANG 8 (fauteuils roulants), cité
# dans les réponses générées — le tri par termes saillants ne le sauve pas non
# plus (il n'est classé qu'en position 8 après tri). Un plafond à 5 l'aurait
# coupé. RÉFUTÉE : le code reste, testé, mais désactivé par défaut. Pistes
# restantes, non essayées : tri par date (résultats post-corpus d'abord),
# `_STOPWORDS_FR` alignée sur `_QE_STOP` (ne plus créditer les mots génériques
# du champ social) — à mesurer avant d'y toucher.
F3_MAX_RESULTATS_INTERNET = None  # None ou 0 = désactivé (défaut depuis le 06/09) ; entier = plafond actif


def filtrer_resultats_internet(results: list, question: str,
                               min_commun: int = 2, min_garde: int = 2,
                               max_retenus: Optional[int] = F3_MAX_RESULTATS_INTERNET) -> tuple:
    """F3 — écarte les résultats internet hors sujet et les doublons ; plafonne
    le reste SI `max_retenus` est activé (entier > 0 — désactivé par défaut,
    `None`/`0`, depuis que la mesure du 06/09 a réfuté le plafond à 5 sur #5621).

    - doublons : même extrait normalisé (boilerplate « 8 nouveaux sites
      universitaires » répété) → un seul conservé ;
    - pertinence : nombre de termes saillants (≥ 4 lettres, hors mots vides)
      communs entre la question et titre + extrait ;
    - retenus, triés par ce score décroissant (à score égal : ordre du moteur) :
      d'abord ceux à ≥ `min_commun` termes communs, plafonnés à `max_retenus`
      si activé ; si moins de `min_garde` au total, on complète avec les
      meilleurs à 1 terme commun, toujours sans dépasser le plafond quand il
      est actif ; un résultat sans aucun terme commun n'est jamais retenu —
      sauf si AUCUN résultat n'en a (la question n'a donné aucun terme
      exploitable : on garde les 3 premiers, ou moins si le plafond actif l'exige).

    Renvoie (retenus dans l'ordre d'origine du moteur, nombre écartés).
    """
    results = list(results or [])
    q_tokens = _tokens_saillants(question)
    vus, uniques = set(), []
    for r in results:
        k = _cle_extrait(r)
        if not k or k in vus:
            continue
        vus.add(k)
        uniques.append(r)
    # plafond effectif : max_retenus si activé (entier > 0), sinon "pas de
    # plafond" — matérialisé par le nombre total de résultats uniques.
    plafond = max_retenus if max_retenus else len(uniques)
    scores = [len(q_tokens & _tokens_saillants(f"{r.get('title', '')} {r.get('content') or r.get('snippet') or ''}"))
              for r in uniques]
    if not any(scores):
        kept = uniques[:min(3, plafond)]
    else:
        # tri stable par score décroissant : à score égal, l'ordre du moteur
        # (celui de `uniques`) est préservé — pas de nouveau signal introduit.
        par_score = sorted(range(len(uniques)), key=lambda i: -scores[i])
        forts = [uniques[i] for i in par_score if scores[i] >= min_commun]
        kept = forts[:plafond]
        if len(kept) < min_garde:
            place = plafond - len(kept)
            faibles = [uniques[i] for i in par_score if 0 < scores[i] < min_commun]
            kept += faibles[:max(0, min(min_garde - len(kept), place))]
        kept = [r for r in uniques if r in kept]  # ordre d'origine du moteur
    return kept, len(results) - len(kept)


def formater_recherche_internet(results: list, max_tokens_total: int,
                                max_results: int = 15, answer: str = "") -> str:
    """Bloc RECHERCHE INTERNET du prompt : un paragraphe daté/sourcé par résultat.

    Lit `content` (Tavily, Google normalisé) ou `snippet` (brut) — l'ancien code
    lisait `snippet` sur les résultats Google normalisés en `content` : le bloc
    internet arrivait VIDE dans le prompt avec le moteur Google.
    """
    kept = [r for r in (results or [])[:max_results]
            if (r.get("content") or r.get("snippet") or "").strip()]
    per = max(40, max_tokens_total // max(1, len(kept)))
    entries = [
        f"Résultat {i} {annoter_resultat_internet(r)}\n"
        f"{truncate_text((r.get('content') or r.get('snippet') or '').strip(), max_tokens=per)}"
        for i, r in enumerate(kept, 1)
    ]
    if answer and answer.strip():
        entries.insert(0, answer.strip())
    return "\n\n".join(entries) if entries else "Aucun résultat internet exploitable."


def rerank_parliamentary_by_recency(
    docs: List["ResponseDocument"],
    now: Optional[datetime] = None,
) -> List["ResponseDocument"]:
    now = now or datetime.now()

    def combined_score(doc) -> float:
        d = safe_parse_date(getattr(doc, "date_reponse", None))
        if d == datetime.min:
            age_years = 5.0  # date inconnue : pénalité modérée
        else:
            age_years = max(0.0, (now - d).days / 365.25)
        recency = 0.5 ** (age_years / RECENCY_HALF_LIFE_YEARS)
        similarity = getattr(doc, "score", 0.0) or 0.0
        return similarity * recency

    return sorted(docs, key=combined_score, reverse=True)

# Fonction qui extrait un ordre numérique à partir d'un label en chiffres romains
ROMAN_MAP = {
    "I":1,"II":2,"III":3,"IV":4,"V":5,"VI":6,"VII":7,"VIII":8,"IX":9,"X":10,
    "XI":11,"XII":12,"XIII":13,"XIV":14,"XV":15,"XVI":16,"XVII":17,"XVIII":18,"XIX":19,"XX":20
}
def extract_order(label: str) -> int:
    if not label:
        return float("inf")
    m_roman = re.search(r"\b(Ier|[IVXLCDM]+)\b", label)
    if m_roman:
        return 1 if m_roman.group(1) == "Ier" else ROMAN_MAP.get(m_roman.group(1), float("inf"))
    m_num = re.search(r"\b(\d+)\b", label)
    if m_num:
        return int(m_num.group(1))
    return float("inf")

# --- F7 : situer chaque article dans son code pour borner sa portée ---
# Avant : « Article L4321-14: <titre> » + contenu. Le modèle prenait l'article de
# tête comme pertinent par construction et lui prêtait une portée générale
# (« obligations déontologiques des professionnels de santé » pour un article du
# chapitre des masseurs-kinésithérapeutes, #8922 ; « logements ET ERP » pour un
# article sur les locaux d'habitation, #2891). On injecte maintenant le code et
# la position (livre / titre / chapitre / section) avant le contenu.
_NOMS_CODES = {
    "CASF": "code de l'action sociale et des familles",
    "Code_de_la_santé_publique": "code de la santé publique",
    "Code_de_la_sécurité_sociale": "code de la sécurité sociale",
    "Code_du_travail": "code du travail",
}


def _position_dans_le_code(contexte_hierarchique: str, niveaux: int = 3, max_len: int = 90) -> str:
    """Derniers niveaux de la hiérarchie (hors « Partie … »), abrégés."""
    parts = [p.strip() for p in (contexte_hierarchique or "").split(">") if p.strip()]
    parts = [p for p in parts if not p.lower().startswith("partie")]
    parts = parts[-niveaux:]
    return " > ".join(p if len(p) <= max_len else p[:max_len - 1].rstrip() + "…" for p in parts)


def formater_article_pour_prompt(art: dict, max_tokens: int) -> str:
    """Bloc d'un article dans TEXTES JURIDIQUES APPLICABLES :

        Article L4321-14 — code de la santé publique (Livre III : … > Chapitre Ier : Masseurs-kinésithérapeutes)
        <titre>
        <contenu tronqué>
    """
    code_brut = art.get("collection") or art.get("code") or ""
    code = _NOMS_CODES.get(code_brut, code_brut).strip()
    position = _position_dans_le_code(art.get("contexte_hierarchique", ""))
    situation = code.lower() if code.lower().startswith("code") else code
    if position:
        situation = f"{situation} ({position})" if situation else f"({position})"
    entete = f"Article {art.get('num', 'N/A')}" + (f" — {situation}" if situation else "")
    titre = (art.get("titre") or "").strip()
    contenu = truncate_text(art.get("contenu", "") or "", max_tokens=max_tokens)
    return "\n".join(p for p in (entete, titre, contenu) if p)


# Fonction qui tri les articles pour que la limite de tokens des prompts s'applique intelligemment
def sort_articles_for_prompt(articles: List[dict]) -> List[dict]:
    """
    Trie les articles pour le prompt en priorisant :
    1. Les articles initiaux (provenance="initial") par score décroissant.
    2. Les articles enrichis associés à chaque article initial, dans l'ordre des articles initiaux.
    """
    # 1. Séparer : articles cités dans la question (F2) > initiaux > enrichis
    cited_articles = [art for art in articles if art.get("provenance") == "cited"]
    initial_articles = [art for art in articles if art.get("provenance") == "initial"]
    enriched_articles = [art for art in articles
                         if art.get("provenance") not in ("initial", "cited")]

    # 2. Trier les articles initiaux par score décroissant
    initial_articles_sorted = sorted(
        initial_articles,
        key=lambda x: x.get("score", 0.0) if x.get("score") is not None else 0.0,
        reverse=True
    )

    # 3. Cités d'abord (ils répondent frontalement à la question), puis initiaux, puis enrichis
    return cited_articles + initial_articles_sorted + enriched_articles

# Construit un prompt pour Mistral Large afin de générer une analyse juridique.
def build_legal_analysis_prompt(question: str, articles: List[dict], stats: dict) -> str:
    """
    Construit un prompt pour Mistral Large afin de générer une analyse juridique.
    Les articles sont des dicts normalisés avec les clés :
    num, titre, contenu, contexte_hierarchique, collection, provenance, score.
    """
    max_tokens = 12000  # Limite pour Small/Medium (16k total - 4k pour la réponse)
    articles_context = []
    total_tokens = 0

    # Utilisation directe de la liste d'articles déjà triée
    for art in articles:
        uid = art.get("num", "N/A")
        titre = art.get("titre", f"Article {uid}")
        contenu = art.get("contenu", "")

        # Estimation des tokens pour cet article
        article_text = f"### Article {uid}: {titre}\n{contenu}\n"
        article_tokens = estimate_tokens(article_text)

        if total_tokens + article_tokens > max_tokens:
            # Tronquer le contenu pour rester dans la limite
            allowed_tokens = max_tokens - total_tokens
            truncated = truncate_text(contenu, max_tokens=allowed_tokens)
            articles_context.append(f"### Article {uid}: {titre}\n{truncated}\n[Texte tronqué]\n")
            break
        else:
            articles_context.append(article_text)
            total_tokens += article_tokens

    articles_str = "\n".join(articles_context)

    # Consignes pour Mistral
    prompt = f"""
    [INST]
    **Consignes strictes pour une analyse juridique complète :**
    1. **Structure obligatoire** à respecter impérativement :
       - Introduction (50-100 mots) : rappel du contexte juridique de la question
       - Analyse détaillée (80-90% du contenu) :
         * L'analyse doit être faite exclusivement à partir des articles suivants :
         * {articles_str}
         * Présentation des principes généraux avec citations précises d'un maximum d'articles
         * Présentation des enjeux
       - Conclusion synthétique (100-150 mots)
       - Toute réponse qui se termine par une phrase tronquée sera considérée comme irrecevable

    2. **Exigences de complétude** :
       - Toute phrase doit être grammaticalement complète
       - La dernière phrase doit impérativement résumer un point clé
       - Si le développement dépasse la limite, prioriser :
         1. Les principes fondamentaux
         2. Les exceptions majeures
         3. Un exemple concret d'application

    3. **Style requis** :
       - Style très concis
       - Phrases courtes (20-25 mots max)
       - Un paragraphe = une idée juridique précise
       - Citations systématiques des articles (ex: "L'article R241-12 précise que...")
       - Éviter les formules vagues ("certains cas", "parfois") → préciser les conditions

    **Question à analyser** :
    {question}

    [/INST]
    """
    return prompt

# Appelle l'API Mistral Large pour générer une analyse juridique
def call_mistral_legal_analysis(
    prompt: str,
    max_tokens: int = 12000,  # Valeur par défaut pour Small/Medium
    temperature: float = 0.3,
    model_size: str = "small"  # "large", "medium", "small"
):
    mistral_api_url = "https://api.mistral.ai/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {MISTRAL_API_KEY}",
        "Content-Type": "application/json"
    }

    # Construire le nom du modèle dynamiquement
    model_name = f"mistral-{model_size}-latest"

    # Ajustement automatique de max_tokens en fonction du modèle
    if model_size == "large":
        max_tokens = min(max_tokens, 28000)  # 32k - 4k (marge pour la réponse)
    else:
        max_tokens = min(max_tokens, 12000)  # 16k - 4k (marge pour la réponse)

    # Vérification de la taille du prompt
    prompt_tokens = estimate_tokens(prompt)
    if prompt_tokens > 30000:  # Limite absolue pour Mistral Large
        raise ValueError(f"Prompt trop long: {prompt_tokens} tokens (limite: 30000)")

    def is_complete(response_text):
        if not response_text:
            return False
        last_char = response_text[-1]
        last_sentence = response_text.split('.')[-1].strip()
        return (
            (last_char in ('.', '!', '?')) and
            (len(last_sentence.split()) > 3) and
            (not last_sentence.endswith((':', ';', ','))) and
            (not response_text.endswith(('...', '–', '—')))
        )

    def complete_response(truncated_response):
        completion_prompt = f"""
        [INST]
        Terminez cette analyse juridique de manière complète et professionnelle.
        Analyse en cours: "{truncated_response[-500:]}"

        Consignes strictes:
        1. Résumez en 1 phrase le point juridique en cours
        2. Ajoutez 2-3 phrases de conclusion qui:
           - Synthétisent les points clés
           - Proposent une application pratique
           - Se terminent impérativement par un point
        3. Utilisez un style formel: "En conséquence...", "Ainsi, il ressort que...", etc.
        [/INST]
        """
        try:
            completion_response = requests.post(
                mistral_api_url,
                headers=headers,
                json={
                    "model": model_name,
                    "messages": [{"role": "user", "content": completion_prompt}],
                    "max_tokens": 500,
                    "temperature": 0.2,
                },
                timeout=30
            )
            completion_response.raise_for_status()
            completion = completion_response.json()["choices"][0]["message"]["content"].strip()
            return truncated_response + "\n\n" + completion
        except Exception as e:
            return truncated_response + f"\n\n[Note: La conclusion de cette analyse a été synthétisée. Erreur technique: {str(e)}]"

    try:
        response = requests.post(
            mistral_api_url,
            headers=headers,
            json={
                "model": model_name,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": temperature,
                "max_tokens": max_tokens,
                "top_p": 0.9,
                "frequency_penalty": 0.1,
                "presence_penalty": 0.1
            },
            timeout=120
        )
        response.raise_for_status()
        data = response.json()
        analysis = data["choices"][0]["message"]["content"].strip()

        if not is_complete(analysis):
            analysis = complete_response(analysis)

        return analysis

    except requests.exceptions.RequestException as e:
        return f"Erreur lors de l'appel à l'API Mistral: {str(e)}"

# Fonction de génération d'analyse juridique
def generate_legal_analysis(
    question: str,
    must_contain: str = "",
    max_articles: int = 5,
    threshold: float = 0.5,
    model_size: str = "small",
    search_button: bool = False,
    generate_analysis_button: bool = False
) -> dict:
    """
    Fonction de génération d'analyse juridique.
    """
    captured_output = None
    try:
        import sys
        from io import StringIO
        old_stdout = sys.stdout
        sys.stdout = captured_output = StringIO()

        # --- Étape 1 : Recherche des articles (avec cache) ---
        if "last_articles_search" not in st.session_state or search_button:
            st.session_state.last_articles_search = search_articles(
                query=question,
                limit=max_articles,
                must_contain=must_contain if must_contain else None,
                debug=True,
                threshold=threshold
            )
        articles = st.session_state.last_articles_search

        # --- Étape 2 : Construction de l'arbre législatif ---
        enrichis = build_legislative_tree(articles["sources"])
        stats = enrichis.pop('stats', {}) if isinstance(enrichis, dict) else {}

        # --- Étape 3 : Formatage des sources pour l'affichage ---
        sources = []
        for uid, data in enrichis.items():
            if isinstance(data, dict) and "article" in data:
                art = data["article"]
                num = art.get("num")
                type_code = num[0] if isinstance(num, str) and num else "?"
                sources.append({
                    "uid": uid,
                    "num": num,
                    "titre": art.get("titre"),
                    "contenu": art.get("contenu"),
                    "collection": art.get("collection"),
                    "contexte_hierarchique": art.get("contexte_hierarchique"),
                    "provenance": art.get("provenance", "initial"),
                    "type": type_code,
                    "parents": data.get("parents", []),
                    "enfants": data.get("enfants", []),
                    "score": art.get("score")
                })

        # --- Étape 4 : Fallback si aucun enrichissement ---
        if not sources:
            sources = []
            for art in articles["sources"]:
                num = art.get("num")
                type_code = num[0] if isinstance(num, str) and num else "?"
                sources.append({
                    "uid": num,
                    "num": num,
                    "titre": art.get("titre"),
                    "contenu": art.get("contenu"),
                    "collection": art.get("collection"),
                    "contexte_hierarchique": art.get("contexte_hierarchique"),
                    "provenance": art.get("provenance", "initial"),
                    "type": type_code,
                    "parents": [],
                    "enfants": [],
                    "score": art.get("score")
                })

        # --- Étape 5 : Gestion des cas d'erreur ---
        if not sources:
            st.warning("Aucun article juridique valide après traitement.")
            return {
                "response": "Aucun article juridique valide après traitement.",
                "sources": [],
                "similar_documents": [],
                "debug_logs": captured_output.getvalue() if captured_output else "",
                "stats": stats
            }

        # --- Étape 6 : Mode "Recherche" ---
        if search_button:
            return {
                "response": "Voir les articles dans l'onglet sources.",
                "sources": sources,
                "similar_documents": [],
                "debug_logs": captured_output.getvalue() if captured_output else "",
                "stats": stats
            }

        # --- Étape 7 : Mode "Génération" ---
        elif generate_analysis_button:
            prep_placeholder = st.empty()
            prep_placeholder.markdown("🔧 Production de l'analyse en cours...", unsafe_allow_html=True)

            # Tri des articles pour le prompt (sans modifier l'affichage)
            articles_for_prompt = sort_articles_for_prompt(sources)

            # Construction du prompt avec les articles triés
            prompt = build_legal_analysis_prompt(
                question,
                articles_for_prompt,  # Liste triée pour le prompt
                stats
            )

            # Ajustement de max_tokens selon le modèle
            max_response_tokens = 28000 if model_size == "large" else 12000

            legal_analysis = call_mistral_legal_analysis(
                prompt,
                max_tokens=max_response_tokens,
                temperature=0.3,
                model_size=model_size
            )

            prep_placeholder.empty()

            return {
                "response": legal_analysis,
                "sources": sources,  # Liste originale pour l'affichage
                "similar_documents": [],
                "debug_logs": captured_output.getvalue() if captured_output else "",
                "stats": stats
            }

    except Exception as e:
        return {
            "response": f"Erreur lors de la génération de l'analyse juridique : {str(e)}",
            "sources": [],
            "similar_documents": [],
            "debug_logs": captured_output.getvalue() if captured_output else "",
            "stats": {}
        }
    finally:
        if 'old_stdout' in locals():
            sys.stdout = old_stdout

# Message `system` de la réponse parlementaire : rôle, registre, garde-fous
# d'exactitude et glossaire — partie stable, réutilisée à chaque appel.
SYSTEME_REPONSE_PARLEMENTAIRE = """Vous êtes rédacteur au sein d'un cabinet ministériel. Vous rédigez le projet de réponse à une question écrite (QE) d'un parlementaire, portant sur la sphère sociale et médico-sociale française. Le texte sera publié au Journal officiel.

REGISTRE
- Style administratif, formel, factuel, à la troisième personne (« Le Gouvernement… », « les services de l'État… »). Jamais de première personne, jamais de formule commerciale ni de politesse.
- Prose continue uniquement : aucun titre, aucune puce, aucune liste, aucune numérotation. Les éléments multiples s'enchaînent en phrases complètes reliées par des connecteurs (« par ailleurs », « en outre », « à cet égard »).
- Aucune redondance : ne répétez pas une idée déjà exprimée.
- Le texte rendu est définitif : aucun crochet, aucun emplacement à compléter (« [montant] », « [à préciser] », « [mesures existantes] ») ne doit y subsister. Les crochets des consignes désignent un élément à remplacer par un fait tiré des contextes fournis ; s'il manque, la phrase est reformulée sans lui.
- Terminez par une phrase de clôture réaffirmant l'engagement du Gouvernement, sans élargir à des sujets éloignés.

EXACTITUDE (impératif — une erreur factuelle disqualifie la réponse)
- Ne citez un chiffre, un montant, un effectif, un pourcentage ou une date d'entrée en vigueur QUE s'il figure explicitement dans l'un des contextes fournis autres que le contexte parlementaire (voir ci-dessous). À défaut, restez qualitatif (« un montant revalorisé chaque année », « plusieurs centaines de structures ») : n'avancez jamais une valeur non sourcée, même plausible.
- N'écrivez jamais « loi n° AAAA-NNN », « décret n° AAAA-NNN », ni une référence d'article précise, si ce numéro exact ne figure pas dans un contexte fourni. Désignez alors le texte par son objet (« la loi relative à… », « le décret encadrant… », « un décret d'application est attendu »).
- Ne développez jamais un sigle absent des contextes fournis et du glossaire ci-dessous ; dans le doute, conservez le sigle seul.
- N'affirmez un lien juridique (« l'article X impose Y ») que si le texte fourni l'énonce clairement.
- La portée d'un article est bornée par sa position dans le code, indiquée après son numéro (livre, titre, chapitre, section : profession, public ou type de locaux concernés). Ne l'étendez jamais au-delà : un article du chapitre des masseurs-kinésithérapeutes ne dit rien des autres professions de santé ; un article sur les locaux d'habitation ne dit rien des établissements recevant du public. Si aucun texte fourni ne recoupe le sujet de la question, n'en citez aucun et renvoyez au cadre général.
- En cas de valeurs contradictoires pour une même donnée, retenez celle de la source la plus récente et précisez sa date.
- Le CONTEXTE PARLEMENTAIRE (réponses ministérielles antérieures) sert d'abord au registre, au ton et aux axes d'argumentation. Par défaut, n'en reprenez aucun chiffre, montant, effectif, pourcentage, date d'entrée en vigueur ni numéro de loi, de décret ou d'article : ces éléments y sont datés, donc présumés périmés — ils doivent alors provenir des TEXTES JURIDIQUES APPLICABLES, des DOCUMENTS DE RÉFÉRENCE ou de la RECHERCHE INTERNET.
- EXCEPTION : lorsqu'une réponse ministérielle est marquée « TRAME FACTUELLE ADMISSIBLE » dans le contexte (proximité élevée, récente, même législature) et qu'elle traite le même texte ou le même dispositif que la question, vous pouvez en reprendre l'enchaînement des faits — succession des textes, décisions de justice, dates-clés, échéances — à trois conditions : (1) recouper chaque élément avec les TEXTES JURIDIQUES, les DOCUMENTS DE RÉFÉRENCE ou la RECHERCHE INTERNET ; (2) ne rien reprendre qu'une source plus récente contredit ; (3) ne pas présenter comme actuel un chiffre ou un état du droit que rien de récent ne confirme. Une réponse marquée « registre seulement » (ancienne ou autre législature) reste cantonnée au ton et aux axes.
- Un texte (loi, décret, arrêté) que la RECHERCHE INTERNET ou un DOCUMENT DE RÉFÉRENCE donne comme **publié** — date de parution, référence au Journal officiel, numéro attribué, entrée marquée « TEXTE PUBLIÉ » — est publié : n'écrivez pas qu'il est « en préparation », « à venir », « à l'étude » ou que « le Gouvernement travaille à son élaboration ».
- Privilégiez systématiquement les textes et la stratégie les plus récents.
- Si un document de référence est une « fiche de référence » (chiffres-clés datés, texte récent), c'est la source à privilégier pour tout chiffre, date, numéro de texte ou définition de sigle : une valeur de fiche remplace toute valeur différente trouvée ailleurs — y compris dans le contexte parlementaire ou un rapport — alors réputée périmée. Si une fiche indique de ne pas citer une valeur, ne la citez pas.

GLOSSAIRE (n'introduire ces notions que si elles figurent dans la question ou un contexte)
AJPA = allocation journalière du proche aidant ; APA = allocation personnalisée d'autonomie ; AVA = assurance vieillesse des aidants ; PCH = prestation de compensation du handicap ; MDPH = maison départementale des personnes handicapées ; MDA = maison départementale de l'autonomie ; CNSA = Caisse nationale de solidarité pour l'autonomie ; PFR = plateforme d'accompagnement et de répit ; GIR = groupe iso-ressources ; CMI = carte mobilité inclusion ; RQTH = reconnaissance de la qualité de travailleur handicapé ; ESMS = établissements et services sociaux et médico-sociaux ; IGAS = Inspection générale des affaires sociales ; DREES = direction de la recherche, des études, de l'évaluation et des statistiques."""

_MOIS_FR = ("", "janvier", "février", "mars", "avril", "mai", "juin", "juillet",
            "août", "septembre", "octobre", "novembre", "décembre")

_DETAIL_JURIDIQUE = {
    1: "Aucune référence juridique n'est requise.",
    2: "Une seule phrase mentionne le cadre juridique général (ex. « Conformément au code de la sécurité sociale, … »).",
    3: "Un paragraphe court (2-3 phrases) expose le cadre juridique applicable, en citant un article clé si pertinent.",
    4: "Un paragraphe (4-5 phrases) analyse les implications juridiques, en citant explicitement deux à trois articles ou principes.",
    5: "Une analyse complète (un à deux paragraphes) cite précisément tous les articles pertinents, leurs interactions et leurs implications concrètes pour la question posée.",
}

_LONGUEUR_MOTS = {
    "Courte": "250 à 350 mots",
    "Moyenne": "450 à 600 mots",
    "Longue": "900 à 1200 mots",
}


# Construit le message `user` de l'appel à Mistral (le `system` est
# SYSTEME_REPONSE_PARLEMENTAIRE, ajouté par call_mistral_parliamentary_response).
def build_parlementary_response_prompt(
    question: str,
    parliamentary_context: str,
    legal_context: str,  # Chaîne de caractères déjà triée et tronquée
    uploaded_documents: str,
    detail_juridique: int,
    longueur: str,
    response_orientation: str,
    custom_instructions: str,
    search_context: str,
    subquestions: Optional[List[str]] = None,
) -> str:
    """
    Construit le message `user` : question, contexte récupéré (parlementaire,
    juridique, documentaire, web), demandes à traiter, et consignes propres à
    la demande (orientation, détail juridique, longueur, instructions libres).
    """
    # Mapping des orientations de réponse
    orientation_mapping = {
        "Répondre de façon neutre":
            "Adoptez un ton neutre et factuel. "
            "Commencez par **souligner l'importance du sujet** pour le Gouvernement, sans reprendre les termes critiques du parlementaire. "
            "Utilisez des formulations comme : "
            "'Ce sujet est une priorité pour le Gouvernement, comme en témoignent [mesures existantes]', "
            "'Le Gouvernement est pleinement conscient des enjeux liés à [thème]', "
            "'Cette question, essentielle pour [public concerné], fait l'objet d'une attention constante de la part des services de l'État'. "
            "Évitez absolument les formulations du type : 'comme vous le soulignez à juste titre', 'vous avez raison de pointer', ou 'la situation est effectivement préoccupante'. "
            "Privilégiez les faits, les chiffres, et les actions en cours.",

        "Répondre négativement aux propositions du parlementaire":
            "Répondez de manière **polie mais ferme**, en **recentrant le débat sur les actions du Gouvernement** plutôt que sur les critiques. "
            "Structurez votre réponse ainsi : "
            "1. **Reconnaissez l'importance du sujet** (sans valider les critiques) : "
            "'La question que vous soulevez touche à un enjeu majeur pour [public concerné], auquel le Gouvernement apporte une réponse structurée.' "
            "2. **Rappelez le cadre existant** : "
            "'Conformément à [texte juridique ou politique publique], les actions menées visent à [objectif].' "
            "3. **Expliquez les contraintes** (si nécessaire) : "
            "'Les marges de manœuvre sont encadrées par [contrainte légale/budgétaire], mais le Gouvernement agit dans le respect de ces règles pour [objectif].' "
            "4. **Mettez en avant les alternatives ou mesures en cours** : "
            "'Plutôt que [proposition du parlementaire], le Gouvernement a choisi de [mesure alternative], qui permet de [bénéfice].' "
            "Exemple : 'Plutôt qu'une refonte complète du dispositif, nous avons renforcé [mesure X], qui a déjà permis [résultat].' "
            "Évitez les formulations défensives comme 'nous ne pouvons pas' – préférez 'notre approche privilégie [solution], car [raison].'",

        "Répondre positivement aux propositions du parlementaire":
            "Saluez l'intérêt de la proposition **sans reprendre les critiques sous-jacentes**. "
            "Utilisez des formulations comme : "
            "'Votre proposition s'inscrit dans une dynamique que le Gouvernement partage, comme en attestent [mesures existantes].' "
            "'Nous partageons votre préoccupation pour [enjeu], et nos actions vont dans le sens de [objectif], comme le montre [exemple].' "
            "Évitez : 'Vous avez raison de souligner que...' → préférez : 'Votre attention à ce sujet rejoint nos priorités, illustrées par [action].'",

        "Répondre de manière technique et détaillée":
            "Fournissez une réponse **factuelle et technique**, en évitant tout commentaire sur les critiques du parlementaire. "
            "Structurez ainsi : "
            "1. **Cadre juridique** : 'Le dispositif actuel, défini par [article X], repose sur [principe].' "
            "2. **Données chiffrées** : 'Les derniers chiffres (source : [DREES/INSEE/...], [année]) montrent que [tendance].' "
            "3. **Mesures en cours** : 'Pour répondre à ces enjeux, [mesure A] et [mesure B] ont été mises en place, avec [résultat].' "
            "Utilisez un vocabulaire neutre et des verbes d'action : 'le Gouvernement a engagé', 'les services travaillent à', 'les résultats montrent que'."
    }

    now = datetime.now(pytz.timezone("Europe/Paris"))
    date_du_jour = f"{now.day} {_MOIS_FR[now.month]} {now.year}"

    detail = detail_juridique if detail_juridique in _DETAIL_JURIDIQUE else 3
    detail_consigne = _DETAIL_JURIDIQUE[detail]
    if detail >= 2:
        detail_consigne += (" Ne citez le numéro d'un article que s'il figure dans les "
                            "textes juridiques fournis ci-dessus ; sinon, renvoyez au code "
                            "concerné sans numéro.")

    mots_cible = next((v for k, v in _LONGUEUR_MOTS.items() if longueur.startswith(k)),
                      "450 à 600 mots")

    if subquestions:
        demandes = "\n".join(f"{i}. {sq}" for i, sq in enumerate(subquestions, 1))
    else:
        demandes = "Déduire les demandes de la question elle-même."

    custom_line = (
        f"- Instruction spécifique impérative : {custom_instructions}\n"
        if custom_instructions else ""
    )

    prompt = f"""Nous sommes le {date_du_jour}.

QUESTION DU PARLEMENTAIRE
{question}

DEMANDES À TRAITER — traitez chacune, dans cet ordre. Si le contexte fourni ne permet pas de répondre précisément à une demande, indiquez-le (renvoi au cadre général, mesure à l'étude, temporisation) — n'avancez jamais un chiffre, une date ou un texte non étayé pour combler.
{demandes}

CONTEXTE PARLEMENTAIRE — réponses ministérielles antérieures à des questions proches. Registre, ton et axes d'argumentation. Une entrée marquée « TRAME FACTUELLE ADMISSIBLE » (proximité élevée, récente, même législature) et portant sur le même texte / dispositif que la question peut fournir l'enchaînement des faits, à recouper avec les autres contextes ; une entrée « registre seulement » n'apporte que le ton.
{parliamentary_context}

TEXTES JURIDIQUES APPLICABLES — prioritaires en cas de contradiction avec une autre source. Chaque article est situé dans son code (livre, titre, chapitre, section) : cette position borne sa portée. Un article qui ne recoupe pas le sujet de la question n'est pas à citer.
{legal_context}

DOCUMENTS DE RÉFÉRENCE
{uploaded_documents}

RECHERCHE INTERNET — actualités et positions du Gouvernement. Chaque résultat est daté et sourcé. Une entrée marquée « TEXTE PUBLIÉ » atteste que le texte visé est paru : tenez-le pour publié (et applicable à sa date d'entrée en vigueur), jamais pour « à venir » ou « en préparation ».
{search_context}

CONSIGNES DE RÉDACTION
- Orientation : {orientation_mapping.get(response_orientation, "")}
- Les crochets [ … ] des tournures ci-dessus sont des emplacements : remplacez-les par un élément des contextes fournis, ou reformulez la phrase sans eux. Aucun crochet ne doit subsister dans la réponse.
- Avant de rédiger, repérez dans les contextes fournis les éléments qui répondent DIRECTEMENT à la question : texte applicable et sa succession, décision de justice, échéance, chiffre daté, position récente du Gouvernement. Construisez la réponse dessus. Ne restez pas au niveau général quand un contexte contient une réponse précise ; à l'inverse, si aucun contexte ne traite frontalement une demande, dites-le (sujet à l'étude, non tranché) sans meubler.
- Corps de la réponse : rappel du cadre juridique et des chiffres disponibles, puis mesures en cours en privilégiant les plus récentes et l'année budgétaire courante. Intégrez les éléments des documents de référence puis de la recherche internet sans nommer la source.
- N'annoncez pas d'échéance déjà passée à la date du jour.
- Si une proposition du parlementaire est pertinente, indiquez qu'elle sera étudiée.
- Niveau de détail juridique : {detail}/5. {detail_consigne}
- Longueur cible : {mots_cible}. Si le sujet ne tient pas dans cette limite, concentrez-vous sur l'enjeu, le cadre juridique et la conclusion. Ne rendez jamais une réponse qui s'achève sur une phrase tronquée.
{custom_line}
PROJET DE RÉPONSE :"""
    return prompt

# Placeholders laissés en clair (« [montant non précisé] ») : compté, pas retiré —
# la boucle d'éval mesure l'effet du prompt, un nettoyage post-hoc le masquerait.
_PLACEHOLDER_RE = re.compile(r"\[[^\]\n]{1,80}\]")


def compter_placeholders(text: str) -> int:
    """Nombre de fragments entre crochets laissés dans une réponse générée."""
    return len(_PLACEHOLDER_RE.findall(text or ""))


# Appelle l'API Mistral Large pour générer une réponse parlementaire
# --- F10 : fin de réponse dédoublée / tronquée en plein mot ---
# Diagnostic (cas #3171, #8922, 05/09) : `finish_reason == "length"` ne veut PAS
# dire que le texte est visiblement incomplet — le modèle atterrit parfois sur
# une phrase de clôture propre pile au niveau du plafond de tokens. L'ancien
# code déclenchait quand même `handle_truncated_response()` sur la seule
# position du dernier point (« < 100 caractères de la fin »), sans vérifier si
# la réponse se terminait déjà correctement. Conséquence : on demandait à un
# second appel de « compléter » un texte déjà fini -> le modèle en tire une
# suite thématique inventée, sur un sujet différent (#3171 : dérive vers les
# infirmiers en pratique avancée, hors sujet IBODE) ou une reprise partielle
# de la dernière phrase (#8922). Cette suite parasite était en plus coupée
# en plein mot (`max_tokens=60` sur l'appel de complétion, alors que sa propre
# consigne demande jusqu'à 400 caractères ≈ 100-130 tokens) — d'où le fragment
# orphelin (« vise à optim. », « "pte. »).
_PONCTUATION_FIN = (".", "!", "?", "»", '"', ")")


def _se_termine_proprement(text: str) -> bool:
    """Vrai si le texte se termine par une ponctuation de fin de phrase (ou une
    fermeture usuelle juste après). Décide si une réponse `finish_reason ==
    "length"` a néanmoins atterri sur une fin propre — auquel cas il ne faut
    PAS tenter de la « compléter » (F10)."""
    t = (text or "").rstrip()
    return bool(t) and t[-1] in _PONCTUATION_FIN


def _completion_utilisable(completion: str, tronquee: bool) -> str:
    """F10 — filet de sécurité sur l'appel de complétion lui-même : s'il a été
    coupé en plein mot (`tronquee=True` ET ne se termine pas proprement), on le
    ramène à sa dernière phrase complète plutôt que de laisser un mot tronqué.
    Chaîne vide si aucune phrase complète n'y figure (rien d'exploitable à
    ajouter — mieux vaut ne rien ajouter qu'un fragment cassé)."""
    if not tronquee or _se_termine_proprement(completion):
        return completion
    dernier_point = completion.rfind(".")
    return completion[:dernier_point + 1] if dernier_point > 0 else ""


def call_mistral_parliamentary_response(
    prompt: str,
    longueur: str,
    question: str,
    max_retries: int = 2,
    model_size: str = "small"
) -> str:
    mistral_api_url = "https://api.mistral.ai/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {MISTRAL_API_KEY}",
        "Content-Type": "application/json"
    }

    model_name = f"mistral-{model_size}-latest"
    max_tokens = 500 if longueur.startswith("Courte") else 1000 if longueur.startswith("Moyenne") else 2200

    # Vérifier la taille du prompt (système + utilisateur)
    prompt_tokens = estimate_tokens(SYSTEME_REPONSE_PARLEMENTAIRE + prompt)
    max_allowed_prompt_tokens = 32000 if model_size == "large" else 16000

    if prompt_tokens > max_allowed_prompt_tokens:
        raise ValueError(f"Prompt trop long: {prompt_tokens} tokens (limite: {max_allowed_prompt_tokens})")

    payload = {
        "model": model_name,
        "messages": [
            {"role": "system", "content": SYSTEME_REPONSE_PARLEMENTAIRE},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.3,
        "max_tokens": max_tokens
    }

    for attempt in range(max_retries):
        try:
            response = requests.post(mistral_api_url, headers=headers, json=payload, timeout=90)
            response.raise_for_status()
            data = response.json()

            mistral_response = data["choices"][0]["message"]["content"]
            is_truncated = data["choices"][0].get("finish_reason") == "length"

            # F10 : ne tenter la complétion que si le texte ne se termine pas
            # déjà proprement — sinon on demanderait de « compléter » une
            # réponse finie, ce qui produit une suite parasite hors sujet.
            if is_truncated and not _se_termine_proprement(mistral_response):
                try:
                    mistral_response = handle_truncated_response(mistral_response, question, longueur)
                except Exception as e:
                    raise Exception(f"Erreur lors de la complétion de la réponse tronquée: {str(e)}")

            return mistral_response

        except requests.exceptions.HTTPError as e:
            if response.status_code == 429 and attempt < max_retries - 1:
                st.warning(f"⏳ API Mistral temporairement encombrée. Tentative {attempt + 1}/{max_retries}. Relance dans 10 secondes...")
                time.sleep(10)
            else:
                raise Exception(f"Erreur HTTP {response.status_code}: {str(e)}")

        except requests.exceptions.RequestException as e:
            if attempt < max_retries - 1:
                st.warning(f"⚠️ Erreur réseau. Tentative {attempt + 1}/{max_retries}. Relance dans 10 secondes...")
                time.sleep(10)
            else:
                raise Exception(f"Erreur réseau: {str(e)}")

# Complète une réponse tronquée (F10 : appelée seulement quand la réponse ne
# se termine PAS proprement, cf. _se_termine_proprement — voir le point d'appel)
def handle_truncated_response(response: str, question: str, longueur: str) -> str:
    # Le point d'appel garantit que `response` est déjà réellement incomplète :
    # tout ce qui suit le dernier point est donc le fragment tronqué (F10 —
    # plus besoin de deviner via une position arbitraire dans le texte).
    last_period = response.rfind('.')
    incomplete_part = response[last_period + 1:] if last_period > 0 else response[-150:]

    completion_prompt = f"""Complétez UNIQUEMENT la phrase ou le paragraphe en cours, sans titre ni introduction, sans recommencer la réponse.

Sujet de la question : {question[:200]}...
Texte à compléter : "{incomplete_part}"

Consignes :
- Continuez directement le texte existant et terminez la phrase ou le paragraphe de manière cohérente.
- Ajoutez au plus 1 à 2 phrases de conclusion réaffirmant l'engagement du Gouvernement.
- Style administratif et formel, prose continue (aucun titre, aucune puce, aucune liste), comme au Journal officiel.
- 400 caractères maximum. Pas de formule de politesse. Ne terminez pas sur une phrase tronquée."""

    try:
        completion_response = requests.post(
            "https://api.mistral.ai/v1/chat/completions",
            headers={"Authorization": f"Bearer {MISTRAL_API_KEY}", "Content-Type": "application/json"},
            json={
                "model": "mistral-large-latest",
                "messages": [{"role": "user", "content": completion_prompt}],
                "temperature": 0.1,
                # F10 : 60 tokens (~240 c) contredisaient la consigne « 400
                # caractères maximum » ci-dessus (≈ 100-130 tokens en français)
                # -> la complétion elle-même se faisait couper en plein mot.
                "max_tokens": 150
            },
            timeout=30
        )
        completion_response.raise_for_status()
        completion_data = completion_response.json()
        completion = completion_data["choices"][0]["message"]["content"]
        completion_truncated = completion_data["choices"][0].get("finish_reason") == "length"
        # F10 : filet de sécurité si la complétion, malgré le budget relevé,
        # est quand même coupée en plein mot — on la ramène à sa dernière
        # phrase complète plutôt que de laisser un fragment cassé.
        completion = _completion_utilisable(completion, completion_truncated)

        if not completion:
            final_response = response
        elif response.endswith("..."):
            final_response = response[:-3] + completion
        elif response.endswith(" "):
            final_response = response + completion
        else:
            final_response = response + " " + completion

        if not _se_termine_proprement(final_response):
            final_response += "."
        return final_response

    except Exception as e:
        st.warning(f"⚠️ Impossible de compléter la réponse tronquée: {str(e)}")
        return response + " (réponse incomplète)"

# --- 7. Génération de la réponse ---

# Fonction qui identifie les sous-questions
def extract_subquestions(question: str) -> list[str]:
    """
    Décompose une question parlementaire en sous-questions explicites ou implicites.
    - Maximum 5 sous-questions
    - Minimum: seulement celles qui sont pertinentes
    - Chaque sous-question doit être formulée clairement
    """
    url = "https://api.mistral.ai/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {MISTRAL_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "mistral-small-latest",
        "messages": [
            {
                "role": "system",
                "content": (
                    "Extrayez les grandes catégories de sous-questions (explicites et implicites) d'une question parlementaire en français. "
                    "NE VOUS APPUYEZ PAS exclusivement sur les points d'interrogation. "
                    "Repèrez les formulations comme: « souhaite savoir », « souhaite donc savoir », « demande », "
                    "« interroge », « voudrait connaître », « comment », « quelles mesures », « envisage », "
                    "« avec quels moyens », « de quelle manière », « clarifier ». "
                    "Pour chaque sous-question ou catégorie de sous-questions : "
                    "- Reformulez-la de manière claire et concise (une vingtaine de mots) en une phrase interrogative. "
                    "- N'inventez rien : si tu n'identifies aucune sous-question, réponds « Aucune sous-question explicite ou implicite détectée. ». "
                    "Développez systématiquement et dans toutes les sous-questions les acronymes en toutes lettres"
                    "**Consigne stricte** : les sous-questions ne doivent pas être redondantes donc si la thématique est très proche il faut les regrouper. "
                    "Format strict, jusqu'à 5 sous-questions ou groupes de sous-questions, maximum mais ça peut être moins : "
                    "1. [Question reformulée] \n"
                    "2. [Question reformulée] \n"
                    "Ne fournissez aucune introduction ni conclusion."
                )
            },
            {
                "role": "user",
                "content": (
                    "Texte à analyser :\n---BEGIN TEXT---\n"
                    f"{question}\n---END TEXT---\n"
                    "Liste TOUTES les questions (explicites ou implicites) posées par le parlementaire."
                )
            }
        ],
        "temperature": 0.0,
        "max_tokens": 350
    }


    response = requests.post(url, headers=headers, json=payload)
    response.raise_for_status()
    data = response.json()

    try:
        text = data["choices"][0]["message"]["content"].strip()
    except (KeyError, IndexError):
        return []

    # Découper en lignes
    subquestions = [line.strip("-• ").strip() for line in text.split("\n") if line.strip()]
    return subquestions[:5]

# Fonction qui génère la réponse à une question parlementaire
def generate_response(
    question: str,
    legislature: Optional[str] = None,
    rubrique: Optional[str] = None,
    detail_juridique: int = 1,
    longueur: str = "Courte (300 mots)",
    response_orientation: str = "Répondre de façon neutre",
    custom_instructions: str = "",
    include_legal_articles: bool = False,
    must_contain: str = "",
    max_legal_articles: int = 3,
    model_size: str = "small"
):
    try:
        status_placeholder = st.empty()

        # Limites de tokens pour Mistral Large
        TOKEN_LIMITS = {
            "large": {
                "system_prompt": 2000,
                "question": 1500,
                "parliamentary_context": 4500,  # 3 réponses QE max
                "search_context": 3000,         # 15 snippets Google
                "uploaded_documents": 5000,     # 10 chunks max
                "legal_context": 5000,          # À affiner plus tard
                "response": 6000,
                "safety_margin": 3000,
            },
            "medium": {
                "system_prompt": 1500,
                "question": 1000,
                "parliamentary_context": 3000,  # 2 réponses QE max
                "search_context": 1500,         # 10 snippets Google
                "uploaded_documents": 2500,     # 5 chunks max
                "legal_context": 2000,          # À affiner plus tard
                "response": 4000,               # Réponse plus longue que small
                "safety_margin": 1500,
            },
            "small": {
                "system_prompt": 1500,          # Même valeur que medium
                "question": 1000,               # Même valeur que medium
                "parliamentary_context": 3000,  # 2 réponses QE max (même que medium)
                "search_context": 1500,         # 10 snippets Google (même que medium)
                "uploaded_documents": 2500,     # 5 chunks max (même que medium)
                "legal_context": 2000,          # Même que medium
                "response": 2000,               # Réponse plus courte que medium
                "safety_margin": 1500,          # Même que medium
            }
        }

        # Utilisation directe des paramètres reçus
        current_orientation = response_orientation
        current_longueur = longueur
        current_detail = detail_juridique
        current_include_legal = include_legal_articles
        current_model_size = model_size

        # Sous-questions explicitées : calculées ICI (une seule fois) pour pouvoir
        # distiller la requête fiche (F9, étape 3) ; réutilisées telles quelles
        # pour le bloc DEMANDES à l'étape 5 (aucun appel Mistral supplémentaire —
        # avant, ce même appel était refait à l'étape 5).
        try:
            subquestions_prompt = extract_subquestions(question)
        except Exception:
            subquestions_prompt = []

        # Étape 1 : Recherche d'anciennes questions
        parliamentary_context = "Aucun contexte parlementaire trouvé."
        similar_documents = []
        if not (hasattr(st.session_state, 'use_priority_docs') and
                st.session_state.use_priority_docs and
                hasattr(st.session_state, 'selected_docs') and
                st.session_state.selected_docs):
            status_placeholder.markdown(
                '<div class="status-message">🏛️ Recherche dans la base des anciennes questions / réponses...</div>',
                unsafe_allow_html=True
            )
            _pool = search_question_parlementaire(question, top_k=8)
            similar_documents = _pool[:5]  # affichage : les 5 plus proches
            if _pool:
                # Contexte du prompt : 3 QE après repondération par la récence
                # (une réponse ancienne ne doit pas dicter la réponse actuelle).
                selected_docs = rerank_parliamentary_by_recency(_pool)[:3]
                _now = datetime.now()
                parliamentary_context = "\n\n".join(
                    [f"Contexte parlementaire {i+1} {annoter_qe_contexte(doc, _now)}:\n"
                     f"Question: {doc.question}\nRéponse: {truncate_text(doc.reponse, max_tokens=TOKEN_LIMITS[current_model_size]['parliamentary_context'] // 3)}"
                     for i, doc in enumerate(selected_docs)]
                )

        # Étape 2 : Recherche des articles juridiques
        legal_context = "Aucun texte juridique spécifique n'a été identifié."
        legal_sources = []
        if (not (hasattr(st.session_state, 'use_priority_docs') and
                st.session_state.use_priority_docs and
                hasattr(st.session_state, 'selected_docs') and
                st.session_state.selected_docs)) and current_include_legal:
            status_placeholder.markdown(
                '<div class="status-message">📚 Recherche dans les codes juridiques...</div>',
                unsafe_allow_html=True
            )
            legal_sources_result = search_articles(
                query=question,
                partie=None,
                limit=max_legal_articles,
                must_contain=must_contain if must_contain else None,
                debug=True,
                threshold=0.5
            )
            legal_sources = legal_sources_result["sources"]

            if legal_sources:
                legal_sources_for_prompt = sort_articles_for_prompt(legal_sources)
                # Tronquer chaque article juridique
                # F7 : chaque article est situé dans son code (portée bornée)
                _par_article = TOKEN_LIMITS[current_model_size]['legal_context'] // len(legal_sources_for_prompt)
                legal_context = "\n\n".join(
                    formater_article_pour_prompt(art, _par_article)
                    for art in legal_sources_for_prompt
                )

        # Étape 3 : Recherche dans les documents uploadés
        status_placeholder.markdown(
            '<div class="status-message">📄 Recherche dans la base documentaire...</div>',
            unsafe_allow_html=True
        )
        # F9 : requête fiche distillée, mode débrayable via DISTILLER_REQUETE_FICHE
        # ("off"/"sujet"/"sujet_sousquestions") — la recherche dans les autres
        # collections-documents continue sur la question brute (fiche_query
        # n'affecte que les collections fiche_de_reference*, cf.
        # search_uploaded_documents).
        fiche_query = None
        if DISTILLER_REQUETE_FICHE != "off":
            try:
                if DISTILLER_REQUETE_FICHE == "resume":
                    fiche_query = resumer_requete_fiche(question)
                else:
                    _subqs_pour_fiche = (subquestions_prompt
                                         if DISTILLER_REQUETE_FICHE == "sujet_sousquestions"
                                         else None)
                    fiche_query = distiller_requete_fiche(question, _subqs_pour_fiche)
            except Exception:
                fiche_query = None
        uploaded_results = search_uploaded_documents(question, qdrant_client, embedding_model, top_k=10, fiche_query=fiche_query)  # + extraits de fiches de référence
        uploaded_docs_context = format_uploaded_docs_by_relevance(uploaded_results, min_score=SEUIL_DOC_STANDARD, max_docs=16)

        # Étape 4 : Recherche internet
        search_context = "Aucune recherche internet effectuée."
        search_results = []

        search_engine = st.session_state.get("search_engine")
        if (search_engine and
            not (hasattr(st.session_state, 'use_priority_docs') and
                 st.session_state.use_priority_docs and
                 hasattr(st.session_state, 'selected_docs') and
                 st.session_state.selected_docs)):
            status_placeholder.markdown(
                f'<div class="status-message">🌐 Recherche internet ({search_engine})...</div>',
                unsafe_allow_html=True
            )

            if search_engine == "Tavily":
                results = search_tavily_government(extract_subject(question))
            else:  # Google
                results = search_google_government(extract_subject(question))
            # F3 : doublons et résultats hors sujet écartés (la liste filtrée est
            # aussi celle affichée dans l'onglet — le modèle et l'utilisateur voient la même chose)
            search_results, _ecartes = filtrer_resultats_internet(results.get("results", []), question)
            # Bloc daté / sourcé, ≤ 15 résultats, marqueur « TEXTE PUBLIÉ » (F1)
            search_context = formater_recherche_internet(
                search_results,
                TOKEN_LIMITS[current_model_size]['search_context'],
                answer=results.get("answer", ""),
            )
            if _ecartes and len(search_results) < 3:
                search_context += (f"\n\n(Peu de résultats pertinents : {_ecartes} résultat(s) "
                                   "écarté(s) comme hors sujet ou en doublon.)")

        # Étape 5 : Génération de la réponse
        status_placeholder.markdown(
            '<div class="status-message">🤖 Génération de la réponse par Mistral...</div>',
            unsafe_allow_html=True
        )

        # subquestions_prompt calculé à l'étape 1 (F9) — réutilisé tel quel ici.

        prompt = build_parlementary_response_prompt(
            question=truncate_text(question, max_tokens=TOKEN_LIMITS[current_model_size]['question']),
            parliamentary_context=parliamentary_context,
            legal_context=legal_context,
            uploaded_documents=uploaded_docs_context,
            detail_juridique=current_detail,
            longueur=current_longueur,
            response_orientation=current_orientation,
            custom_instructions=custom_instructions,
            search_context=search_context,
            subquestions=subquestions_prompt,
        )

        mistral_response = call_mistral_parliamentary_response(
            prompt,
            current_longueur,
            question,
            model_size=current_model_size
        )

        status_placeholder.empty()

        # Préparation des métadonnées
        metadata = {
            "mode": "parlementaire",
            "timestamp": datetime.now(pytz.timezone('Europe/Paris')).isoformat(),
            "model_used": f"mistral-{current_model_size}-latest",
            "include_legal_articles": current_include_legal,
            "longueur": current_longueur,
            "response_orientation": current_orientation,
            "detail_juridique": current_detail,
            "legislature": legislature,
            "rubrique": rubrique,
            "custom_instructions": custom_instructions,
            "placeholders_restants": compter_placeholders(mistral_response),
        }

        return {
            "question": question,
            "context": [doc.reponse for doc in similar_documents[:6] if hasattr(doc, 'reponse')],
            "context_str": parliamentary_context,
            "legal_context": legal_context,
            "response": mistral_response,
            "legal_sources": legal_sources if current_include_legal else [],
            "similar_documents": similar_documents,  # Retourner les 5 documents similaires
            "uploaded_documents": uploaded_results,
            "search_results": search_results,
            "metadata": metadata
        }

    except Exception as e:
        return {
            "question": question,
            "response": f"Erreur lors de la génération de la réponse : {str(e)}",
            "error": str(e),
            "legal_sources": [],
            "similar_documents": [],
            "metadata": {
                "status": "error",
                "timestamp": datetime.now(pytz.timezone('Europe/Paris')).isoformat(),
                "legislature": legislature,
                "rubrique": rubrique
            }
        }

# --- NOUVEL ENDPOINT SIMPLIFIÉ POUR LISTER LES DOCUMENTS ---

# Fonction qui retourne juste les deux catégories de documents - VERSION STATIQUE
def get_simple_documents_list():
    return {
        "documents": [
            {
                "type": "Questions écrites (QE) de l'Assemblée nationale",
                "periode": "2017-2025",
                "description": "Questions ayant obtenu une réponse ministérielle (avant le 1er novembre 2025)."
            },
            {
                "type": "Questions écrites (QE) du Sénat",
                "periode": "2017-2025",
                "description": "Collection complète des questions écrites ayant obtenu une réponse ministérielle (avant le 1er novembre 2025)."
            }
        ]
    }

#################################################################
### -------------  2. AUTHENTIFICATION -----------------------###
#################################################################

config = {
    "credentials": {
        "usernames": {
            "Whisler": {
                "name": "Francois-Mathieu",
                "password": os.getenv("USER_WHISLER_PASSWORD")
            },
            "Delphine": {
                "name": "Caudilla Delphine",
                "password": os.getenv("USER_DELPHINE_PASSWORD")
            },
            "Isabelle": {
                "name": "Caudilla Isabelle",
                "password": os.getenv("USER_ISABELLE_PASSWORD")
            },
            "Arnaud": {
                "name": "Arnaud",
                "password": os.getenv("USER_ARNAUD_PASSWORD")
            },
            "DGCS": {
                "name": "DGCS",
                "password": os.getenv("USER_DGCS_PASSWORD")
            },
            "DSS": {
                "name": "DSS",
                "password": os.getenv("USER_DSS_PASSWORD")
            },
            "Julien": {
                "name": "Special Guest",
                "password": os.getenv("USER_SPECIAL_PASSWORD")
            },
            "Invité": {
                "name": "Invité",
                "password": os.getenv("USER_GUEST_PASSWORD")
            }
        }
    }
}

if 'authentication_status' not in st.session_state:
    st.session_state.authentication_status = None

def check_password():
    if st.session_state["username"] in config['credentials']['usernames']:
        stored_password = config['credentials']['usernames'][st.session_state["username"]]["password"]
        if hashlib.sha256(st.session_state["password"].encode()).hexdigest() == stored_password:
            st.session_state["authentication_status"] = True
            st.session_state["name"] = config['credentials']['usernames'][st.session_state["username"]]["name"]
            return
    st.session_state["authentication_status"] = False

# --- Connexion ou contenu principal ---

if st.session_state.authentication_status is not True:
    # Masquer la sidebar avant connexion
    st.markdown("""
        <style>
        [data-testid="stSidebar"] {display: none;}
        /* Centrer le titre */
        .auth-title {text-align: center; margin-top: 0.5rem;}
        /* Centrer le paragraphe d'intro */
        .auth-intro {text-align: center; color: #5c5c5c;}
        </style>
    """, unsafe_allow_html=True)

    # --- Page d'authentification ---
    st.markdown('<h1 class="auth-title">🔐 Authentification requise</h1>', unsafe_allow_html=True)
    st.markdown('<p class="auth-intro">Veuillez entrer vos identifiants pour accéder au générateur de réponses aux questions écrites.</p>', unsafe_allow_html=True)

    # Colonnes pour réduire la largeur et centrer les champs
    # Ajuste les ratios pour obtenir la largeur souhaitée (ici ~25% de la page)
    left, center, right = st.columns([3, 2, 3])
    with center:
        st.text_input("Nom d'utilisateur", key="username")
        st.text_input("Mot de passe", type="password", key="password")
        if st.button("Se connecter"):
            check_password()
            st.rerun()

    if st.session_state.authentication_status is False:
        st.error("Identifiants incorrects. Veuillez réessayer.")

else:
    # Réafficher la sidebar après connexion

    st.markdown("""
        <style>
        [data-testid="stSidebar"] {display: block;}
        </style>
    """, unsafe_allow_html=True)

#################################################################
### ---- 3. AFFICHAGE DU SITE APRES CONNEXION --------------- ###
#################################################################

    # --- Initialisation de l'historique (UNIQUEMENT si non existant) ---
    if "full_historique" not in st.session_state:
        st.session_state.full_historique = {}
        save_historique()  # Sauvegarde initiale

    # --- Chargement de l'historique depuis session_state (si vide, on vérifie le cache) ---
    if not st.session_state.full_historique and "historique_cache" in st.session_state:
        st.session_state.full_historique = st.session_state.historique_cache

# --- Configuration de la page et CSS ---
    st.markdown("""
    <style>
        /* Élargir le conteneur principal */
        .block-container {
            max-width: 95%;
            padding-left: 2rem;
            padding-right: 2rem;
        }
        /* Élargir les zones de saisie */
        textarea, .stTextArea textarea {
            width: 100% !important;
        }
        /* Élargir les selectbox et sliders */
        .stSelectbox, .stSlider {
            width: 100% !important;
        }
        /* Votre CSS existant */
        .stApp { background-color: #f8f9fa; }
        .stTabs [data-baseweb="tab-list"] { gap: 0; background-color: #e9ecef; border-radius: 6px 6px 0 0; padding: 4px; }
        .stTabs [data-baseweb="tab"] { height: 36px; white-space: pre-wrap; background-color: #f8f9fa; border: none; border-radius: 4px 4px 0 0; padding: 0 12px; }
        .stTabs [aria-selected="true"] { background-color: #ffffff; font-weight: bold; color: #3d3d3d; }
        .stButton>button { background-color: #4a8bfc; color: white; border: none; border-radius: 4px; padding: 8px 16px; font-weight: 500; }
        .stButton>button:hover { background-color: #3a7bfc; }
        .stExpander { background-color: #ffffff; border: 1px solid #e9ecef; border-radius: 6px; margin-bottom: 0px; } /* ← réduit l’espace */
        .stTextArea textarea { font-family: 'Segoe UI', sans-serif; font-size: 16px; line-height: 1.5; }
        .stAlert { border-radius: 6px; }
        .source-text { font-family: monospace; font-size: 14px; background-color: #f8f9fa; padding: 8px; border-radius: 4px; border-left: 3px solid #4a8bfc; }
        .response-text { font-family: 'Segoe UI', sans-serif; font-size: 16px; line-height: 1.6; white-space: pre-wrap; background-color: white; padding: 16px; border-radius: 6px; border: 1px solid #e9ecef; }

        /* Nouveau : titre de l'historique */
        .history-title {
            margin-bottom: 15px;
        }

        /* Style carte avec ombre légère pour les expanders */
        div[data-testid="stExpander"] {
            background-color: #ffffff;
            border: 1px solid #e9ecef;
            border-radius: 8px;
            margin-bottom: 5px !important;
            box-shadow: 0 2px 6px rgba(0,0,0,0.08); /* ombre douce */
        }

    </style>
    """, unsafe_allow_html=True)

    st.title("🏛️ Générateur de réponses aux questions écrites parlementaires")
    st.markdown("""
    Application de réponse aux questions parlementaires,
    appuyée sur une base documentaire (embedding avec **camemBERT** finetuné), un moteur de recherche (**Tavily** ou **Google**) et les modèles **Mistral**.
    """)

#################################################################
#### -------------------- 3a. SIDEBAR  --------------------- ####
#################################################################

    with st.sidebar:

        # Message de bienvenue + version
        st.write(f'Bienvenue *{st.session_state["name"]}*')
        st.caption("Version v0.1.12 (Beta)")

        st.markdown("---")

        # Choix du module
        mode = st.radio(
            "Choix du module",
            ["Réponse parlementaire", "Base documentaire", "Analyse juridique"], # Ajouter , "Analyse juridique" pour avoir le second mode - Permet de moduler les modes accessibles sur le site
            index=0
        )

        # Moteur de recherche internet : fixé en conf (MOTEUR_RECHERCHE_INTERNET),
        # ce n'est plus un choix exposé à l'utilisateur.
        st.session_state["search_engine"] = MOTEUR_RECHERCHE_INTERNET
        # Modèle Mistral : défaut ; le sélecteur est dans « Options avancées »
        # (chemin parlementaire). Seule source lue à la génération.
        st.session_state.setdefault("model_size", MODELE_MISTRAL_DEFAUT)

        st.markdown("---")

        # Déconnexion (en bas de la sidebar)
        if st.button('Déconnexion', key="logout"):
            st.session_state.authentication_status = None
            st.rerun()

    # Boutons de génération : définis dans la colonne principale, sous chaque
    # formulaire. Initialisés ici pour éviter tout NameError si le mode courant
    # ne les crée pas.
    generate_parliamentary_button = False
    generate_analysis_button = False


##########################################################################
### ------ 3b. INTERFACE DE GESTION DE LA BASE DOCUMENTAIRE ---------- ###
##########################################################################

    # Initialisation du bouton search
    search_button = False
    
    if mode == "Base documentaire":
        st.markdown("---")
        st.markdown("#### 📚 Gestion de la base documentaire")

        # 1. Liste des documents (collections)
        try:
            collections = qdrant_client.get_collections()
            collection_names = [col.name for col in collections.collections]

            protected_collections = {
                "Code de la sécurité sociale",
                "Code du travail",
                "CASF",
                "QuestionParlementaire",
                "Code de la santé publique"
            }

            # Filtre pour ne garder que les collections "documents" (ex: "NomDuDocument_2023")
            doc_collections = [
                name for name in collection_names
                if name not in protected_collections and "_" in name  # Ex: "MonDocument_2023"
                and not _is_legal_infra_collection(name)
            ]

            if not doc_collections:
                st.info("Aucun document trouvé.")
            else:
                st.markdown("**Documents disponibles :**")

                # Tri alphabétique sur le nom nettoyé
                sorted_collections = sorted(
                    doc_collections,
                    key=lambda name: name.split('__')[0].replace('_', ' ').lower()
                )

                for doc_name in sorted_collections:
                    clean_name = doc_name.split('__')[0].replace('_', ' ')
                    col1, col2, col3 = st.columns([18, 1, 1])  # plus de place pour le nom

                    with col1:
                        st.markdown(f"📄 {clean_name}")

                    with col2:
                        if st.button("✏️", key=f"rename_{doc_name}", help="Renommer le document"):
                            st.session_state.show_rename_modal = True
                            st.session_state.current_doc_to_rename = doc_name
                            st.rerun()

                    with col3:
                        if st.button("🗑️", key=f"del_{doc_name}", help="Supprimer le document"):
                            qdrant_client.delete_collection(collection_name=doc_name)
                            st.success(f"Document '{clean_name}' supprimé.")
                            st.rerun()

                # Fenêtre modale de renommage
                if st.session_state.show_rename_modal:
                    doc_name = st.session_state.current_doc_to_rename
                    clean_name = doc_name.split('__')[0].replace('_', ' ')

                    with st.container():
                        st.markdown("---")
                        st.subheader(f"Renommer le document ''{clean_name}''")

                        new_name = st.text_input(
                            "Nouveau nom:",
                            value=clean_name.replace(" ", "_"),
                            key=f"new_name_{doc_name}"
                        )

                        col1, col2 = st.columns(2)

                        with col1:
                            if st.button("Valider", key=f"validate_rename_{doc_name}"):
                                if new_name.strip():
                                    try:
                                        # 1. Vérification si le nouveau nom existe déjà
                                        new_display_name = new_name.strip().replace(" ", "_")
                                        new_collection_name = f"{new_display_name}__{doc_name.split('__')[1]}"

                                        # Récupère toutes les collections existantes
                                        existing_collections = qdrant_client.get_collections()
                                        existing_names = [col.name for col in existing_collections.collections]

                                        # Vérifie si le nouveau nom existe déjà
                                        if new_collection_name in existing_names:
                                            st.error(f"❌ Le nom '{new_name}' existe déjà. Veuillez choisir un autre nom.")
                                        else:
                                            # Initialisation de la progression
                                            progress_bar = st.progress(0)
                                            status_text = st.empty()
                                            estimated_total = 1000  # Estimation conservatrice du nombre total de points

                                            # 2. Crée la nouvelle collection (0-10%)
                                            status_text.text("Création de la nouvelle collection...")
                                            qdrant_client.create_collection(
                                                collection_name=new_collection_name,
                                                vectors_config=models.VectorParams(
                                                    size=1024,
                                                    distance=models.Distance.COSINE
                                                )
                                            )
                                            progress_bar.progress(10)

                                            # 3. Récupère les points (10-40%)
                                            status_text.text("Récupération des données...")
                                            offset = None
                                            all_points = []

                                            while True:
                                                records, offset = qdrant_client.scroll(
                                                    collection_name=doc_name,
                                                    limit=100,
                                                    offset=offset,
                                                    with_payload=True,
                                                    with_vectors=True
                                                )

                                                for record in records:
                                                    point_dict = {
                                                        "id": str(record.id),
                                                        "vector": record.vector,
                                                        "payload": record.payload
                                                    }
                                                    all_points.append(models.PointStruct(**point_dict))

                                                if offset is None:
                                                    break

                                                # Calcul sécurisé de la progression (max 40%)
                                                current_progress = min(40, 10 + int(30 * len(all_points) / max(1, estimated_total)))
                                                progress_bar.progress(current_progress)

                                            # 4. Copie les points (40-90%)
                                            status_text.text(f"Copie des {len(all_points)} chunks...")
                                            batch_size = 50

                                            for i in range(0, len(all_points), batch_size):
                                                batch = all_points[i:i + batch_size]
                                                qdrant_client.upsert(
                                                    collection_name=new_collection_name,
                                                    points=batch,
                                                    wait=True
                                                )

                                                # Calcul SÉCURISÉ de la progression (max 90%)
                                                batch_progress = min(90, 40 + int(50 * (i + len(batch)) / max(1, len(all_points))))
                                                progress_bar.progress(batch_progress)

                                            # 5. Supprime l'ancienne collection (100%)
                                            status_text.text("Finalisation...")
                                            qdrant_client.delete_collection(collection_name=doc_name)
                                            progress_bar.progress(100)

                                            st.success(f"✅ Document renommé en '{new_name}' !")
                                            st.session_state.show_rename_modal = False
                                            st.rerun()

                                    except Exception as e:
                                        st.error(f"Erreur: {e}")
                                else:
                                    st.warning("Veuillez entrer un nom valide.")


                        with col2:
                            if st.button("Annuler", key=f"cancel_rename_{doc_name}"):
                                st.session_state.show_rename_modal = False
                                st.rerun()

        except Exception as e:
            st.error(f"Erreur lors de la récupération des collections : {e}")

        # --- Section de téléchargement depuis un lien public ---
        st.markdown("---")
        st.markdown("**Ajouter un document (max 50 Mo) depuis un lien public**")

        def make_direct_link(file_url: str, prefer_format: str = "docx") -> str:
            """
            Convertit les liens Google Docs / Google Drive / Dropbox / OneDrive en liens directs téléchargeables.
            - prefer_format: "docx" ou "pdf" pour Google Docs.
            Retourne le lien original si aucun cas particulier n'est détecté.
            """
            u = file_url.strip()

            # --- Google Docs (document, pas un fichier Drive) ---
            if "docs.google.com/document" in u:
                m = re.search(r"/document/d/([^/]+)/", u)
                doc_id = m.group(1) if m else None
                if doc_id:
                    fmt = "docx" if prefer_format == "docx" else "pdf"
                    return f"https://docs.google.com/document/d/{doc_id}/export?format={fmt}"
                return u

            # --- Google Drive (fichiers) ---
            if "drive.google.com" in u:
                file_id = None
                if "/d/" in u:
                    file_id = u.split("/d/")[1].split("/")[0]
                else:
                    qs_id = parse_qs(urlparse(u).query).get("id", [None])[0]
                    file_id = qs_id or file_id
                if file_id:
                    return f"https://drive.google.com/uc?export=download&id={file_id}"
                return u

            # --- Dropbox ---
            if "dropbox.com" in u:
                # Force le téléchargement direct
                if "dl=" in u:
                    u = re.sub(r"dl=\d", "dl=1", u)
                else:
                    sep = "&" if urlparse(u).query else "?"
                    u = u + f"{sep}dl=1"
                # Variante possible : remplacer dl=1 par raw=1
                # u = u.replace("dl=1", "raw=1")
                return u

            # --- OneDrive ---
            if "1drv.ms" in u or "onedrive.live.com" in u:
                if "download=" not in u:
                    parsed = urlparse(u)
                    sep = "&" if parsed.query else "?"
                    return u + f"{sep}download=1"
                return u

            return u

        file_url = st.text_input(
            "🔗 Collez ici le lien public vers votre document (Google Drive, Dropbox, OneDrive, etc.) :",
            key="file_url_input",
            placeholder="Ex: https://drive.google.com/uc?export=download&id=..."
        )

        if file_url:
            direct_url = make_direct_link(file_url)
            default_name = os.path.splitext(urlparse(direct_url).path.split('/')[-1])[0]
            custom_name = st.text_input(
                "Nom du document (sera aussi le nom de la collection) :",
                value=default_name,
                key="doc_name_input"
            )

            # Bouton de validation du lien
            if st.button("✅ Valider le lien"):
                try:
                    response = requests.get(direct_url, stream=True, timeout=10)
                    if response.status_code == 200:
                        st.success("Lien valide et accessible ✅")
                    else:
                        st.error(f"❌ Lien inaccessible (status {response.status_code})")
                except Exception as e:
                    st.error(f"❌ Erreur lors de la validation : {e}")

            # Bouton de téléchargement
            if st.button("📥 Télécharger et traiter le document"):
                try:
                    response = requests.get(direct_url, stream=True, timeout=30)
                    response.raise_for_status()

                    file_size = int(response.headers.get('content-length', 0))
                    if file_size > 50 * 1024 * 1024:
                        st.error("❌ Fichier trop gros (max 50 Mo).")
                    else:
                        with st.spinner("Téléchargement en cours..."):
                            downloaded = 0
                            total_size = file_size
                            progress_bar = st.progress(0)

                            # Déterminer l’extension via content-type ou fallback
                            content_type = response.headers.get("content-type", "").lower()

                            if "pdf" in content_type or direct_url.lower().endswith(".pdf"):
                                file_extension = "pdf"
                            elif "docx" in content_type or "wordprocessingml" in content_type or direct_url.lower().endswith(".docx"):
                                file_extension = "docx"
                            else:
                                # fallback par défaut : PDF
                                file_extension = "pdf"

                            # Créer un fichier temporaire
                            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as f:
                                for chunk in response.iter_content(chunk_size=8192):
                                    if chunk:
                                        f.write(chunk)
                                tmp_path = f.name

                            # Vérification rapide du contenu
                            with open(tmp_path, "rb") as f:
                                header = f.read(5)
                                if file_extension == "pdf" and not header.startswith(b"%PDF"):
                                    st.error("❌ Le fichier téléchargé n'est pas un PDF valide.")
                                    os.unlink(tmp_path)
                                    st.stop()
                                if file_extension == "docx":
                                    import zipfile
                                    try:
                                        with zipfile.ZipFile(tmp_path, 'r') as z:
                                            if '[Content_Types].xml' not in z.namelist():
                                                st.error("❌ Le fichier téléchargé n'est pas un DOCX valide.")
                                                os.unlink(tmp_path)
                                                st.stop()
                                    except Exception:
                                        st.error("❌ Le fichier téléchargé n'est pas un DOCX valide.")
                                        os.unlink(tmp_path)
                                        st.stop()

                        # Créer une collection Qdrant avec suivi d'avancement
                        collection_name = f"{custom_name.replace(' ', '_')}__{int(datetime.now().timestamp())}"

                        progress = st.progress(0)
                        status = st.empty()

                        status.info("📂 Initialisation de la collection...")
                        progress.progress(10)

                        try:
                            qdrant_client.create_collection(
                                collection_name=collection_name,
                                vectors_config=models.VectorParams(size=1024, distance=models.Distance.COSINE)
                            )
                            status.success(f"✅ Collection '{collection_name}' créée.")
                            progress.progress(30)

                            # Traiter et indexer le document
                            status.info("📑 Extraction et préparation du texte...")
                            success = process_and_index_document(
                                file_path=tmp_path,
                                file_type=file_extension,
                                collection_name=collection_name,
                                qdrant_client=qdrant_client,
                                embedding_model=embedding_model,
                                progress_callback=lambda current, total, message: (
                                    progress.progress(min(current/total, 1.0)),
                                    status.info(message)
                                )
                            )

                            if success:
                                progress.progress(100)
                                status.success(f"✅ Document ajouté sous le nom '{custom_name}' !")
                                collection_info = qdrant_client.get_collection(collection_name)
 #                                st.info(f"📊 Nombre de vecteurs indexés : {collection_info.vectors_count}")
                            else:
                                status.error("❌ Échec du traitement.")

                        except Exception as e:
                            st.error(f"❌ Erreur lors de la récupération des collections : {e}")

                except Exception as e:
                    st.error(f"❌ Echec du traitement du document {e}")

#########################################################################################
#### ---------- 3c. INTERFACE DE REPONSE AUX QUESTIONS PARLEMENTAIRES -------------- ####
#########################################################################################

    elif mode == "Réponse parlementaire" or mode == "Analyse juridique":
        st.markdown("#### Question parlementaire")
        question = st.text_area(
            "Question parlementaire",   # libellé non vide
            height=150,
            placeholder="Copier ici le texte de la question parlementaire",
            key="question_input",
            label_visibility="collapsed"
        )

        # Si le texte change, on efface les sous-questions précédentes
        if "last_question" not in st.session_state:
            st.session_state["last_question"] = ""

        if st.session_state["question_input"] != st.session_state["last_question"]:
            st.session_state["last_question"] = st.session_state["question_input"]
            st.session_state.pop("subquestions", None)

        # Initialisation des variables des boutons (pour éviter NameError)
        search_button = False

        # --- Boutons conditionnels selon le mode ---

        if mode == "Réponse parlementaire":
            st.markdown("#### Paramètres de réponse")

            # Chemin principal : orientation + longueur, toujours visibles.
            col1, col2 = st.columns([2, 1])
            with col1:
                response_orientation_options = [
                    "Répondre de façon neutre",
                    "Répondre négativement aux propositions du parlementaire",
                    "Répondre positivement aux propositions du parlementaire",
                    "Répondre de manière technique et détaillée"
                ]
                response_orientation = st.selectbox(
                    "Orientation de la réponse",
                    response_orientation_options,
                    index=0,
                    key="response_orientation"
                )

            with col2:
                longueur = st.selectbox(
                    "Longueur de la réponse",
                    ["Courte (300 mots)", "Moyenne (500 mots)", "Longue (1000 mots)"],
                    index=0,  # "Courte" par défaut
                    key="longueur"
                )

            # Valeurs par défaut : les widgets ci-dessous vivent dans l'expander
            # « Options avancées » (toujours rendu, même replié) et écrasent ces
            # valeurs. Elles ne servent que de garde-fou si l'expander était retiré.
            detail_juridique = st.session_state.get("detail_juridique", 3)
            include_legal_articles = True
            must_contain = ""
            custom_instructions = ""

            with st.expander("⚙️ Options avancées"):

                # --- Modèle Mistral ---
                _labels = list(_MODEL_LABELS)
                _cur = st.session_state.get("model_size", MODELE_MISTRAL_DEFAUT)
                _idx = next((i for i, lab in enumerate(_labels)
                             if _MODEL_LABELS[lab] == _cur), 0)
                model_label = st.selectbox(
                    "Modèle Mistral", _labels, index=_idx, key="model_label"
                )
                st.session_state["model_size"] = _MODEL_LABELS[model_label]

                st.markdown("---")

                # --- Recherche dans les codes juridiques (activée par défaut) ---
                include_legal_articles = st.toggle(
                    "Rechercher dans les codes juridiques",
                    value=True,
                    key="include_legal_articles",
                )
                if include_legal_articles:
                    cja, cjb = st.columns([1, 2])
                    with cja:
                        detail_juridique = st.slider(
                            "Niveau de détail juridique (1 = bas, 5 = élevé)",
                            min_value=1, max_value=5, value=3,
                            key="detail_juridique",
                            help="Pilote aussi le nombre d'articles de loi récupérés.",
                        )
                    with cjb:
                        must_contain = st.text_input(
                            "Les articles retenus doivent contenir (optionnel)",
                            key="must_contain_input",
                            placeholder="Ex: allocation, article 123, décret 2020-...",
                        )

                st.markdown("---")

                # --- Instructions de rédaction ---
                custom_instructions = st.text_area(
                    "Instructions générales (max. 300 caractères)",
                    placeholder="Ex: Insister sur l'aspect budgétaire, mentionner le projet de loi X...",
                    height=100,
                    max_chars=300,
                    key="custom_instructions",
                )

                st.caption(
                    "Instructions par sous-question — optionnel. Si le champ reste "
                    "vide, la sous-question est traitée sans consigne spécifique."
                )
                if st.button("➕ Décomposer en sous-questions",
                             help="Découpe la question en sous-questions pour leur donner des consignes distinctes"):
                    st.session_state["subquestions"] = extract_subquestions(
                        st.session_state.get("question_input", "")
                    )

                if "subquestions" in st.session_state:
                    for i, sq in enumerate(st.session_state["subquestions"], start=1):
                        with st.container(border=True):
                            st.markdown(f"**Sous-question {i}.** {sq}")
                            st.text_area(
                                f"Instructions (max. 200 caractères)",
                                key=f"instructions_sq_{i}",
                                max_chars=200,
                                height=80,
                                label_visibility="collapsed",
                                placeholder=(
                                    "Ex : renvoyer au débat parlementaire en cours ; "
                                    "rappeler qu'une concertation est engagée ; "
                                    "refuser la proposition au motif que…"
                                    if i == 1 else ""
                                ),
                            )

                st.markdown("---")

                # --- Limiter les sources ---
                use_priority_docs = st.checkbox(
                    "Limiter la recherche à certains documents",
                    value=False,
                    key="use_priority_docs",
                    help="La réponse n'intègre alors ni les anciennes QE, ni les textes juridiques, ni la recherche internet.",
                )
                if use_priority_docs:
                    collections = qdrant_client.get_collections()
                    doc_collections = [
                        col.name for col in collections.collections
                        if col.name not in {"Code_de_la_sécurité_sociale", "Code_du_travail", "CASF", "QuestionParlementaire", "Code_de_la_santé_publique"}
                        and "_" in col.name
                        and not _is_legal_infra_collection(col.name)
                    ]
                    doc_names = [col.split('__')[0].replace('_', ' ') for col in doc_collections]
                    selected_docs = st.multiselect(
                        "Documents à utiliser",
                        options=doc_names,
                        key="priority_docs",
                        placeholder="Choisir...",
                    )

            # Bouton de génération : sous le formulaire, dans le flux de lecture.
            generate_parliamentary_button = st.button(
                "Générer la réponse",
                type="primary",
                key="generate_parliamentary_button",
            )

            st.markdown(
                "<hr style='border:1px solid #bbb; width:100%;'>",
                unsafe_allow_html=True
            )

        elif mode == "Analyse juridique":
            # --- Paramètres de recherche juridique ---
            st.markdown("#### Paramètres de recherche juridique")
            col1, col2, col3 = st.columns(3)
            with col1:
                must_contain = st.text_input("🔎 Doit contenir (mot ou expression exacte)", key="must_contain_input")
            with col2:
                threshold = st.selectbox(
                    "📊 Seuil de sélection",
                    [0.4, 0.5, 0.6, 0.65, 0.70, 0.75, 0.80, 0.85],
                    index=1,
                    key="threshold_select"
                )
            with col3:
                max_articles = st.number_input(
                    "📌 Nombre maximum d'articles",
                    min_value=1,
                    max_value=30,
                    value=5,
                    key="max_articles_input"
                )

            # --- Boutons : rechercher les articles / générer l'analyse ---
            bcol1, bcol2 = st.columns([1, 1])
            with bcol1:
                search_button = st.button(
                    "Rechercher les articles", type="secondary", key="search_button"
                )
            with bcol2:
                generate_analysis_button = st.button(
                    "Générer l'analyse", type="primary", key="generate_analysis_button"
                )

            st.markdown(
                "<hr style='border:1px solid #bbb; width:100%;'>",
                unsafe_allow_html=True
            )

###################################################################
#### --------------- 4. GENERATION DE LA REPONSE ------------- ####
###################################################################

    if generate_parliamentary_button or search_button or generate_analysis_button:
        if not question.strip():
            st.warning("Veuillez entrer une question.")
        else:
            try:
                debug_logs = ""
                response_data = {}

                # Récupération des valeurs depuis st.session_state
                response_orientation = st.session_state.get('response_orientation', "Répondre de façon neutre")
                longueur = st.session_state.get('longueur', "Courte (300 mots)")
                detail_juridique = st.session_state.get('detail_juridique', 3)
                include_legal_articles = bool(st.session_state.get('include_legal_articles', True))
                must_contain = st.session_state.get('must_contain_input', "")
                model_size = st.session_state.get('model_size', MODELE_MISTRAL_DEFAUT)

                if mode == "Réponse parlementaire":
                    response_data = generate_response(
                        question=question,
                        detail_juridique=detail_juridique,
                        longueur=longueur,
                        response_orientation=response_orientation,
                        custom_instructions=custom_instructions,
                        include_legal_articles=include_legal_articles,
                        must_contain=must_contain if include_legal_articles else "",
                        max_legal_articles=detail_juridique,
                        model_size=model_size
                    )

                elif mode == "Analyse juridique":
                    # generate_legal_analysis capture stdout et n'émet aucun statut
                    # dans l'UI : un spinner évite l'écran figé pendant l'appel.
                    with st.spinner("Recherche des articles et analyse en cours…"):
                        response_data = generate_legal_analysis(
                            question=question,
                            must_contain=must_contain,
                            max_articles=max_articles,
                            threshold=threshold,
                            model_size=model_size,
                            search_button=search_button,
                            generate_analysis_button=generate_analysis_button
                        )

                # Détermination des onglets à afficher
                tabs = []
                if mode == "Réponse parlementaire":
                    tabs = ["📜 Réponse"]
                    if include_legal_articles:
                        tabs.append("⚖️ Articles juridiques")
                    tabs.extend(["🏛️ Anciennes QE", "📰 Recherches actualités", "📄 Base documentaire"])
                elif mode == "Analyse juridique":
                    tabs = ["⚖️ Articles juridiques"]
                    if generate_analysis_button and response_data.get("response"):
                        tabs.insert(0, "📜 Analyse")

                # Création dynamique des onglets
                if tabs:
                    st_tabs = st.tabs(tabs)

                    # Affichage du contenu en fonction des onglets
                    for i, tab in enumerate(st_tabs):
                        with tab:
                            # [Votre code existant pour l'affichage des onglets...]
                            if mode == "Réponse parlementaire":
                                if "📜 Réponse" in tabs[i]:
                                    st.markdown("#### Réponse générée")
                                    st.markdown(response_data["response"])
                                    if response_data.get("debug_logs"):
                                        with st.expander("🐛 Voir les logs de recherche"):
                                            st.text_area("Logs", response_data["debug_logs"], height=200)
                                    if response_data.get("response"):
                                        export_content = build_export_content(
                                            response_data,
                                            mode="parlementaire",
                                            include_legal_articles=include_legal_articles
                                        )
                                        st.download_button(
                                            label="📥 Exporter en TXT",
                                            data=export_content.encode("utf-8"),
                                            file_name="export_reponse_parlementaire.txt",
                                            mime="text/plain",
                                            key=f"export_reponse_{i}"
                                        )
                                    st.markdown('<div style="height: 300px;"></div>', unsafe_allow_html=True)

                                elif "🏛️ Anciennes QE" in tabs[i]:
                                    st.markdown("#### 5 Questions parlementaires les plus similaires de la plus récente à la plus ancienne")
                                    if not response_data.get("similar_documents"):
                                        st.info("Aucune question parlementaire similaire trouvée.")
                                    else:
                                        similar_documents = response_data["similar_documents"]

                                        # Tri par date décroissante
                                        similar_documents = sorted(
                                            similar_documents,
                                            key=lambda d: safe_parse_date(d.date_reponse),
                                            reverse=True
                                        )

                                        # Affichage avec score
                                        for idx, doc in enumerate(similar_documents):
                                            chambre = doc.chambre or ("Assemblée nationale" if str(doc.uid).startswith("QAN") else "Sénat")
                                            date_reponse = doc.date_reponse or "Inconnue"
                                            question_text = doc.question or "Question non disponible"
                                            reponse_text = doc.reponse or "Réponse non disponible"
                                            score = f"{doc.score:.2f}" if doc.score is not None else "N/A"

                                            with st.expander(f"{idx+1}. QE {doc.uid} ({chambre}) - Score de proximité : {score}"):
                                                st.markdown(f"**Date de réponse:** {date_reponse}")
                                                st.markdown(f"**Chambre:** {chambre}")
                                                st.markdown(f"**Question:** {question_text}")
                                                st.markdown(f"**Réponse:** {reponse_text}")

                                    st.markdown('<div style="height: 300px;"></div>', unsafe_allow_html=True)

                                elif "⚖️ Articles juridiques" in tabs[i]:
                                    st.markdown("###### 📚 Articles juridiques pertinents")
                                    legal_sources = response_data.get("legal_sources", [])
                                    if not legal_sources:
                                        st.info("Aucun article juridique trouvé.")
                                    else:
                                        # 1. Regroupement par code juridique
                                        by_code = {}
                                        for art in legal_sources:
                                            code = art.get("collection", "Inconnu")
                                            if code not in by_code:
                                                by_code[code] = []
                                            by_code[code].append(art)

                                        # 2. Affichage par code avec hiérarchie textuelle
                                        for code, articles in by_code.items():
                                            st.markdown(f"##### 📚 {code}")

                                            # Regroupement par contexte hiérarchique complet
                                            by_contexte = {}
                                            for art in articles:
                                                contexte = art.get("contexte_hierarchique", "Sans contexte")
                                                if contexte not in by_contexte:
                                                    by_contexte[contexte] = []
                                                by_contexte[contexte].append(art)

                                            # 3. Affichage de chaque contexte avec ses articles
                                            for contexte, arts in by_contexte.items():
                                                # Affichage du contexte hiérarchique en texte
                                                st.markdown(f"###### {contexte.replace('>', ' > ')}")

                                                # Affichage des articles de ce contexte
                                                for art in arts:
                                                    # Détermination du type et de l'icône
                                                    provenance = art.get("provenance", "initial")
                                                    score = art.get("score")

                                                    if provenance == "initial":
                                                        icon = "🔍"
                                                        if score is not None:
                                                            label = f"Article initial - score: {score:.3f}"
                                                        else:
                                                            label = "Article initial"
                                                    elif provenance == "context":
                                                        icon = "📑"
                                                        label = "Article de contexte"
                                                    elif provenance == "reference":
                                                        icon = "🔗"
                                                        label = "Article référencé"
                                                    else:
                                                        icon = "📄"
                                                        label = "Article"

                                                    # Expander pour chaque article
                                                    with st.expander(f"{icon} {art.get('num', '?')} - {art.get('titre', '')} ({label})"):
                                                        st.markdown(f"**Collection:** {art.get('collection', '')}")
                                                        st.markdown(f"**Numéro:** {art.get('num', '')}")
                                                        if provenance == "initial" and score is not None:
                                                            st.markdown(f"**Score:** {score:.3f}")
                                                        st.markdown(f"**Contexte hiérarchique:** {art.get('contexte_hierarchique', '')}")
                                                        st.markdown(f"**Contenu:**\n\n{art.get('contenu', '')}")

                                        # 4. Légende des icônes
                                        st.markdown("""
                                        **Légende des icônes:**
                                        - 🔍: Article "initial" (issu de la recherche dans les codes juridiques, avec score de pertinence)
                                        - 📑: Article de contexte (même section/paragraphe qu'un article "initial")
                                        - 🔗: Article référencé (référencé par un article initial)
                                        """)


                                elif "📰 Recherches actualités" in tabs[i]:
                                    st.markdown("#### Dernières annonces et actualités gouvernementales")

                                    # Afficher le résumé global si disponible
                                    summary = response_data.get("search_context") or response_data.get("answer")
                                    if summary:
                                        st.markdown(f"**Résumé de la recherche :**\n\n{summary}")

                                    search_results = response_data.get("search_results", [])
                                    if not search_results:
                                        st.info("Aucune actualité trouvée via le moteur de recherche.")
                                    else:
                                        for idx, item in enumerate(search_results):
                                            titre = item.get("title", "Sans titre")
                                            url = item.get("url", "")
                                            # Utiliser "content" pour Tavily, "snippet" pour Google
                                            extrait = item.get("content") or item.get("snippet") or ""
                                            with st.expander(f"{idx+1}. {titre}"):
                                                if url:
                                                    st.markdown(f"[Lien vers la source]({url})")
                                                if extrait:
                                                    st.markdown(extrait)

                                elif "📄 Base documentaire" in tabs[i]:
                                    st.markdown("#### Résultats pertinents dans les documents uploadés")

                                    uploaded_results = response_data.get("uploaded_documents", [])

                                    # F12 : MÊME règle que l'injection dans le prompt
                                    # (`extrait_retenu`, clause fiche comprise). Avant,
                                    # ce filtre était un `>= 0.7` codé en dur sans la
                                    # clause fiche : un extrait de fiche entre 0,45 et
                                    # 0,70 était envoyé au modèle sans jamais s'afficher.
                                    filtered_results = [res for res in uploaded_results
                                                        if extrait_retenu(res)]

                                    if not filtered_results:
                                        st.info("Aucun extrait pertinent trouvé dans les documents uploadés.")
                                    else:
                                        # Regrouper par document
                                        grouped = {}
                                        for res in filtered_results:
                                            doc_name = res["collection"].split('__')[0].replace('_', ' ')
                                            grouped.setdefault(doc_name, []).append(res)

                                        for doc_name, results in grouped.items():
                                            with st.expander(f"📄 {doc_name} ({len(results)} extraits)"):
                                                for idx, res in enumerate(results, start=1):
                                                    score = f"{res['score']:.2f}" if res.get("score") is not None else "N/A"
                                                    text_preview = res["text"]
                                                    title = res.get("title") or "N/A"

                                                    # F12 : distinguer à l'œil un extrait fort
                                                    # d'un extrait retenu par la seule clause fiche
                                                    _marginal = extrait_marginal(res)
                                                    _tag = (" — ⚠️ fiche, sous le seuil standard "
                                                            f"({SEUIL_DOC_STANDARD:.2f}) : retenu par la clause "
                                                            "fiche de référence" if _marginal else "")
                                                    st.markdown(f"**Extrait {idx} (score: {score}){_tag}**")
                                                    st.markdown(text_preview)
                                                    st.markdown(f"**Section :** {title}")
                                                    st.markdown("---")

                                                # Actions sur la collection
                                                col1, col2 = st.columns([1, 1])
                                                target_collection = results[0]["collection"]
                                                with col1:
                                                    if st.button("✏️ Renommer", key=f"rename_{target_collection}"):
                                                        st.session_state.show_rename_modal = True
                                                        st.session_state.current_doc_to_rename = target_collection
                                                        st.rerun()
                                                with col2:
                                                    if st.button("🗑️ Supprimer", key=f"del_{target_collection}"):
                                                        qdrant_client.delete_collection(collection_name=target_collection)
                                                        st.success(f"Document '{doc_name}' supprimé.")
                                                        st.rerun()

                            elif mode == "Analyse juridique":
                                if "📜 Analyse" in tabs[i]:
                                    st.markdown("#### Analyse juridique générée")
                                    st.markdown(response_data["response"])

                                    # Affichage des logs (si disponibles)
                                    if response_data.get("debug_logs"):
                                        with st.expander("🐛 Voir les logs de recherche"):
                                            st.text_area("Logs", response_data["debug_logs"], height=200)

                                    # Bouton d'export
                                    if response_data.get("response"):
                                        export_content = build_export_content(
                                            response_data,
                                            mode="analyse",  # Mode pour l'export
                                            include_legal_articles=False
                                        )
                                        st.download_button(
                                            label="📥 Exporter en TXT",
                                            data=export_content.encode("utf-8"),
                                            file_name=f"analyse_juridique_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",  # Nom de fichier unique
                                            mime="text/plain",
                                            key=f"export_analyse_{i}"
                                        )
                                    st.markdown('<div style="height: 300px;"></div>', unsafe_allow_html=True)

                                elif "⚖️ Articles juridiques" in tabs[i]:
                                    st.markdown("###### 📚 Articles juridiques pertinents")
                                    sources = response_data.get("sources", [])
                                    if not sources:
                                        st.info("Aucun article juridique trouvé.")
                                    else:
                                        # 1. Regroupement par code juridique
                                        by_code = {}
                                        for art in sources:
                                            code = art.get("collection", "Inconnu")
                                            if code not in by_code:
                                                by_code[code] = []
                                            by_code[code].append(art)

                                        # 2. Affichage par code avec hiérarchie textuelle
                                        for code, articles in by_code.items():
                                            st.markdown(f"##### 📚 {code}")

                                            # Regroupement par contexte hiérarchique complet
                                            by_contexte = {}
                                            for art in articles:
                                                contexte = art.get("contexte_hierarchique", "Sans contexte")
                                                if contexte not in by_contexte:
                                                    by_contexte[contexte] = []
                                                by_contexte[contexte].append(art)

                                            # 3. Affichage de chaque contexte avec ses articles
                                            for contexte, arts in by_contexte.items():
                                                # Affichage du contexte hiérarchique en texte
                                                st.markdown(f"###### {contexte.replace('>', ' > ')}")

                                                # Affichage des articles de ce contexte
                                                for art in arts:
                                                    # Détermination du type et de l'icône
                                                    provenance = art.get("provenance", "initial")
                                                    score = art.get("score")

                                                    if provenance == "initial":
                                                        icon = "🔍"
                                                        if score is not None:
                                                            label = f"Article initial - score: {score:.3f}"
                                                        else:
                                                            label = "Article initial"
                                                    elif provenance == "context":
                                                        icon = "📑"
                                                        label = "Article de contexte"
                                                    elif provenance == "reference":
                                                        icon = "🔗"
                                                        label = "Article référencé"
                                                    else:
                                                        icon = "📄"
                                                        label = "Article"

                                                    # Expander pour chaque article
                                                    with st.expander(f"{icon} {art.get('num', '?')} - {art.get('titre', '')} ({label})"):
                                                        st.markdown(f"**Collection:** {art.get('collection', '')}")
                                                        st.markdown(f"**Numéro:** {art.get('num', '')}")
                                                        if provenance == "initial" and score is not None:
                                                            st.markdown(f"**Score:** {score:.3f}")
                                                        st.markdown(f"**Contexte hiérarchique:** {art.get('contexte_hierarchique', '')}")
                                                        st.markdown(f"**Contenu:**\n\n{art.get('contenu', '')}")

                                        # 4. Légende des icônes
                                        st.markdown("""
                                        **Légende des icônes:**
                                        - 🔍: Article initial de la recherche (avec score si disponible)
                                        - 📑: Article de contexte (même section/paragraphe)
                                        - 🔗: Article référencé
                                        """)

                # === NOUVELLE PARTIE POUR L'HISTORIQUE ===
                # Génération de la clé unique pour l'historique
                timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                q_hash = hashlib.md5(f"{question}_{timestamp}".encode()).hexdigest()

                # Préparation des métadonnées
                metadata = {
                    "mode": mode,
                    "timestamp": datetime.now(pytz.timezone('Europe/Paris')).isoformat(),
                    "model_used": f"mistral-{model_size}-latest",
                    "include_legal_articles": include_legal_articles,
                    "longueur": longueur,
                    "response_orientation": response_orientation,
                    "detail_juridique": detail_juridique,
                    "legislature": None,
                    "rubrique": None,
                    "custom_instructions": custom_instructions
                }

                # Ajout des métadonnées à response_data
                response_data["metadata"] = metadata

                # Initialisation de l'historique si nécessaire
                if "full_historique" not in st.session_state:
                    st.session_state.full_historique = {}

                # Ajout à l'historique
                st.session_state.full_historique[q_hash] = {
                    "question": question,
                    "response": response_data.get("response", "Pas de réponse générée"),
                    "similar_documents": response_data.get("similar_documents", []),
                    "legal_sources": response_data.get("legal_sources", []),
                    "sources": response_data.get("sources", []),
                    "search_results": response_data.get("search_results", []),
                    "uploaded_documents": response_data.get("uploaded_documents", []),
                    "metadata": metadata
                }
                save_historique()

            except Exception as e:
                st.error(f"Erreur inattendue: {str(e)}")
                st.exception(e)

###################################################################
#### -------- 5. GESTION ET AFFICHAGE DE L'HISTORIQUE -------- ####
###################################################################

    # Affichage de l'historique complet (votre code existant)
    st.markdown("""
    <style>
        .history-container { max-width: 1000px; margin-left: auto; margin-right: 0; padding-left: 1rem; }
        .history-expander { margin-bottom: 0.5rem !important; border: 1px solid #e9ecef; border-radius: 8px; }
        .metadata-line { font-size: 0.8em; color: #666; margin-bottom: 0.5em; margin-top: 0.5em; }
    </style>
    """, unsafe_allow_html=True)

    with st.container():
        st.markdown('<div class="history-container">', unsafe_allow_html=True)
        if hasattr(st.session_state, 'full_historique') and st.session_state.full_historique:
            nb_entries = len(st.session_state.full_historique)
            st.markdown(
                f'<div style="display:flex;align-items:center;gap:10px;margin-bottom:1rem;">'
                f'<span style="font-size:23px;font-weight:bold;">🗂️ Historique des réponses générées ({nb_entries})</span>'
                f'</div>',
                unsafe_allow_html=True
            )

            # Tri du plus récent au plus ancien
            sorted_historique = sorted(
                st.session_state.full_historique.items(),
                key=lambda x: x[1]["metadata"].get("timestamp", ""),
                reverse=True
            )

            for idx, (q_hash, entry) in enumerate(sorted_historique):
                metadata = entry.get("metadata", {})

                custom_instructions = metadata.get('custom_instructions', '')
                custom_instructions_display = f"Instructions: '{custom_instructions}'" if custom_instructions else "Instructions: Aucune"

                metadata_str = (
                    f"🕒 {metadata.get('timestamp', '')[:16].replace('T', ' ')} | "
                    f"Mode: {metadata.get('mode', 'inconnu')} | "
                    f"Modèle: {metadata.get('model_used', 'inconnu')} | "
                    f"Orientation: {metadata.get('response_orientation', 'inconnu')} | "
                    f"Articles juridiques: {'Oui' if metadata.get('include_legal_articles', False) else 'Non'} | "
                    f"Longueur: {metadata.get('longueur', 'inconnu')} | "
                    f"Détail juridique: {metadata.get('detail_juridique', 'inconnu')}/5 | "
                    f"{custom_instructions_display}"
                )

                st.markdown(f'<div class="metadata-line">{metadata_str}</div>', unsafe_allow_html=True)

                with st.expander(f"{idx+1}. {truncate_text(entry['question'], max_tokens=50)}", expanded=False):
                    st.markdown(f"**📝 Question:**\n{entry.get('question', 'Non disponible')}")
                    st.markdown(f"**💬 Réponse:**")
                    st.markdown(entry.get('response', 'Non disponible'))

                    # Boutons d'action
                    col1, col2 = st.columns([1, 1])
                    with col1:
                        export_content = build_export_content(
                            entry,
                            mode=metadata.get("mode", "parlementaire"),
                            include_legal_articles=metadata.get("include_legal_articles", False)
                        )
                        st.download_button(
                            label="⬇️ Exporter en TXT",
                            data=export_content.encode("utf-8"),
                            file_name=f"export_{q_hash[:8]}_{metadata.get('mode', 'parlementaire')}.txt",
                            mime="text/plain",
                            key=f"export_hist_{q_hash}"
                        )
                    with col2:
                        if st.button("🗑️ Supprimer", key=f"del_hist_{q_hash}"):
                            del st.session_state.full_historique[q_hash]
                            save_historique()
                            st.rerun()

                    # --- Sources juridiques (sans expander) ---
                    if entry.get("legal_sources") or entry.get("sources"):
                        st.markdown("### ⚖️ Articles juridiques")

                        # Fonction pour afficher un article juridique
                        def display_legal_article(article_data):
                            if not isinstance(article_data, dict):
                                st.warning(f"Structure inattendue: {type(article_data)}")
                                return

                            # Clés adaptées à la structure actuelle
                            num = article_data.get("num", "N/A")
                            collection = article_data.get("collection", "N/A")
                            titre = article_data.get("titre", "Non disponible")
                            texte = article_data.get("contenu", "Non disponible")
                            contexte = article_data.get("contexte_hierarchique", "Non disponible")
                            provenance = article_data.get("provenance", "initial")
                            score = article_data.get("score")

                            # Affichage "à plat" (sans expander)
                            st.markdown(f"#### Article {num} ({collection})")
                            st.markdown(f"**Titre:** {titre}")
                            st.markdown(f"**Contexte hiérarchique:** {contexte}")
                            if provenance == "initial" and score is not None:
                                st.markdown(f"**Score:** {score:.3f}")
                            st.markdown(f"**Texte:**\n{texte}")
                            st.markdown("---")

                        # Mode "analyse_juridique" (utilise "sources")
                        if metadata.get("mode") == "analyse_juridique":
                            sources = entry.get("sources", [])
                            if not sources:
                                st.info("Aucun article juridique enregistré.")
                            else:
                                # Regroupement par code juridique
                                by_code = {}
                                for art in sources:
                                    code = art.get("collection", "Inconnu")
                                    if code not in by_code:
                                        by_code[code] = []
                                    by_code[code].append(art)

                                # Affichage par code
                                for code, articles in by_code.items():
                                    st.markdown(f"**Code: {code}**")
                                    for art in articles:
                                        display_legal_article(art)

                        # Mode "Réponse parlementaire" (utilise "legal_sources")
                        else:
                            legal_sources = entry.get("legal_sources", [])
                            if not legal_sources:
                                st.info("Aucun article juridique disponible.")
                            else:
                                for article in legal_sources:
                                    display_legal_article(article)

                    # --- Résultats de recherche internet (sans expander) ---
                    if entry.get("search_results"):
                        st.markdown("### 🌐 Résultats de recherche internet")
                        # Résumé global
                        if entry.get("search_context") or entry.get("answer"):
                            st.markdown(f"**Résumé :** {entry.get('search_context') or entry.get('answer')}")
                        for item_idx, item in enumerate(entry["search_results"]):
                            st.markdown(f"#### Résultat {item_idx+1}: {item.get('title', 'Sans titre')}")
                            if item.get("url"):
                                st.markdown(f"[Lien]({item.get('url')})")
                            st.markdown(item.get("content") or item.get("snippet", "Aucun extrait disponible"))
                            st.markdown("---")

                    # --- Anciennes QE similaires (sans expander) ---
                    if entry.get("similar_documents"):
                        st.markdown("### 🏛️ Questions parlementaires similaires")
                        for doc_idx, doc in enumerate(entry["similar_documents"]):
                            st.markdown(f"#### QE {doc_idx+1} - Score: {getattr(doc, 'score', 'N/A'):.2f}")
                            st.markdown(f"**Question:** {getattr(doc, 'question', 'Non disponible')}")
                            st.markdown(f"**Réponse:** {getattr(doc, 'reponse', 'Non disponible')}")
                            st.markdown("---")

                    # --- Résultats vectoriels sur documents uploadés (sans expander) ---
                    if entry.get("uploaded_documents"):
                        st.markdown("### 📄 Documents uploadés pertinents")
                        grouped = {}
                        for res in entry["uploaded_documents"]:
                            doc_name = res["collection"].split('__')[0].replace('_', ' ')
                            grouped.setdefault(doc_name, []).append(res)
                        for doc_name, results in grouped.items():
                            st.markdown(f"#### {doc_name} ({len(results)} extraits)")
                            for idx_res, res in enumerate(results, start=1):
                                score = f"{res['score']:.2f}" if res.get("score") is not None else "N/A"
                                text_full = res["text"]
                                title = res.get("title") or "N/A"
                                st.markdown(f"**Extrait {idx_res} (score: {score})**")
                                st.markdown(text_full)
                                st.markdown(f"**Section :** {title}")
                                st.markdown("---")

        else:
            st.info("Aucune question enregistrée dans l'historique pour le moment.")
        st.markdown('</div>', unsafe_allow_html=True)
