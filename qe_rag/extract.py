"""Extraction de texte — PDF / DOCX / TXT / HTML (revue § B1).

`app.py::extract_text` ne gère que le PDF (`pdfplumber`), et
`process_and_index_document` que pdf/docx. Une grande partie du corpus visé est
du web (Legifrance, solidarites.gouv, dossiers législatifs AN/Sénat) : d'où le
support HTML ici, avec la cascade trafilatura -> readability -> BeautifulSoup ->
strip regex (dépendances optionnelles, importées à la demande).
"""
from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Union
from urllib.parse import urlparse

_UA = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/126.0 Safari/537.36"
)
_HEADERS = {"User-Agent": _UA, "Accept-Language": "fr-FR,fr;q=0.9"}

_KIND_BY_EXT = {
    ".pdf": "pdf", ".docx": "docx", ".doc": "docx",
    ".txt": "txt", ".md": "txt", ".text": "txt",
    ".html": "html", ".htm": "html",
}


def is_url(s: str) -> bool:
    p = urlparse(s)
    return p.scheme in ("http", "https")


def guess_kind(src: str, explicit: str | None = None) -> str:
    if explicit:
        return explicit
    if is_url(src):
        ext = Path(urlparse(src).path).suffix.lower()
        return _KIND_BY_EXT.get(ext, "html")
    return _KIND_BY_EXT.get(Path(src).suffix.lower(), "")


# --------------------------------------------------------------------------- #
def fetch_bytes(url: str, timeout: int = 90) -> bytes:
    import requests

    r = requests.get(url, headers=_HEADERS, timeout=timeout)
    r.raise_for_status()
    return r.content


# --------------------------------------------------------------------------- #
def extract_pdf(data: Union[bytes, str, Path]) -> str:
    """pdfplumber si dispo, sinon pypdf ; bascule pypdf aussi si pdfplumber rend trop peu."""
    src_bytes = None
    if isinstance(data, (bytes, bytearray)):
        src_bytes = bytes(data)

    text = ""
    try:
        import pdfplumber

        src = io.BytesIO(src_bytes) if src_bytes is not None else str(data)
        parts = []
        with pdfplumber.open(src) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                if t.strip():
                    parts.append(t)
        text = "\n".join(parts)
    except ImportError:
        pass
    except Exception:  # noqa: BLE001
        text = ""

    if len(text) < 300:
        try:
            from pypdf import PdfReader

            src = io.BytesIO(src_bytes) if src_bytes is not None else str(data)
            text2 = "\n".join((p.extract_text() or "") for p in PdfReader(src).pages)
            if len(text2) > len(text):
                text = text2
        except Exception:  # noqa: BLE001
            pass
    return text


def extract_docx(path: Union[str, Path]) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # noqa: BLE001
        raise RuntimeError("DOCX : `pip install python-docx`") from exc
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_txt(data: Union[bytes, str, Path]) -> str:
    raw = bytes(data) if isinstance(data, (bytes, bytearray)) else Path(data).read_bytes()
    for enc in ("utf-8", "cp1252", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def extract_html(data: Union[bytes, str], url: str = "") -> str:
    html = data.decode("utf-8", errors="replace") if isinstance(data, (bytes, bytearray)) else data

    try:
        import trafilatura

        t = trafilatura.extract(
            html, url=url or None, include_comments=False, include_tables=True, favor_recall=True
        )
        if t and len(t) > 400:
            return t
    except Exception:  # noqa: BLE001
        pass

    try:
        from bs4 import BeautifulSoup
        from readability import Document as RDoc

        t = BeautifulSoup(RDoc(html).summary(), "lxml").get_text("\n")
        if t and len(t) > 400:
            return t
    except Exception:  # noqa: BLE001
        pass

    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "nav", "header", "footer", "aside", "form"]):
            tag.decompose()
        cand = max(
            soup.find_all(["article", "main", "div", "section"]),
            key=lambda e: len(e.get_text(" ", strip=True)),
            default=soup.body or soup,
        )
        return cand.get_text("\n") if cand else ""
    except Exception:  # noqa: BLE001
        pass

    # dernier recours : strip brut des balises
    text = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", " ", html)
    text = re.sub(r"(?s)<[^>]+>", "\n", text)
    return re.sub(r"\n{2,}", "\n", text)


# --------------------------------------------------------------------------- #
def extract(src: str, kind: str | None = None) -> tuple[str, str]:
    """`src` = chemin de fichier OU URL. Renvoie (texte_brut, kind_effectif)."""
    kind = guess_kind(src, kind)
    if not kind:
        raise ValueError(f"type de document indéterminé pour {src!r} — préciser kind=")

    if is_url(src):
        data = fetch_bytes(src)
        if kind == "pdf":
            return extract_pdf(data), kind
        if kind == "txt":
            return extract_txt(data), kind
        return extract_html(data, src), "html"

    p = Path(src)
    if not p.exists():
        raise FileNotFoundError(src)
    if kind == "pdf":
        return extract_pdf(p.read_bytes()), kind
    if kind == "docx":
        return extract_docx(p), kind
    if kind == "html":
        return extract_html(p.read_bytes(), ""), kind
    return extract_txt(p), "txt"
